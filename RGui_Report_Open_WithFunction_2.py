import os
import pandas as pd
from PyQt5.QtWidgets import QDialog
from PyQt5 import QtWidgets
from Screening_System_PyQt5 import Screening_Report_WordGenerator
from Screening_System_PyQt5 import RawData_Report_WordDoc
from Screening_System_PyQt5 import GraphPicture_Report_Combin
from PyQt5.uic import loadUi
from docx import Document
import sys
import _thread
from docx import Document
from docxcompose.composer import Composer

pd.options.display.max_columns = 999
pd.options.display.max_rows = 999
pd.set_option("display.precision", 6)


class Report_Open_WithFunctions(QDialog):
    def __init__(self):
        super(Report_Open_WithFunctions, self).__init__()
        self.dir_path = os.path.dirname(sys.argv[0])

    def combine_word_documents(self, files, testname):

        # print(f"files= {files}, testname={testname}")
        # path_name = os.path.dirname(sys.argv[0]) + r"\\Report_Word\\"
        # master = Document(path_name + testname + '.docx')
        # print(path_name + testname + '.docx')
        # composer = Composer(master)
        # for item in files:
        #     print(path_name + item)
        #     doc_add = Document(path_name + item)
        #     print("this is doc_add")
        #     print(doc_add)
        #     composer.append(doc_add)
        #     print("this get print or not: composer.append")
        #     print(composer.append(doc_add))
        #
        # composer.save(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\" + testname + '.docx')

        # old working version
        merged_document = Document()
        files.insert(0, testname+'.docx')
        print(files)
        # it was the post-static report giving the error
        # files.pop(3)
        # print(files)

        for index, file in enumerate(files):
            sub_doc = Document(os.path.dirname(sys.argv[0]) + r"\\Report_Word\\" + file)

            # if index < len(files) - 1:
            #     sub_doc.add_page_break()

            for element in sub_doc.element.body:
                merged_document.element.body.append(element)

        merged_document.save(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\" + testname + '.docx')
        #
        _thread.start_new_thread(os.system, (os.path.dirname(sys.argv[0]) + r"\\Final_Report\\" + testname + '_Graph.docx',))
        _thread.start_new_thread(os.system, (os.path.dirname(sys.argv[0]) + r"\\Final_Report\\" + testname + '.docx',))

    def OK_Pressed(self, selected_file):
        # self.listWidget.currentItem().text() = 14872B00.txt
        Screening_Report_WordGenerator.screening_Report_wordGenerator(selected_file)
        # the statistic table report is generated when gui is generate the code are together
        RawData_Report_WordDoc.RawData_report_wordDoc(selected_file)
        GraphPicture_Report_Combin.graphPicture_combin(selected_file)
        testName = selected_file[0:-4]
        files = []
        files.append("Note-"+testName + ".docx")
        files.append(testName + "Pre-StaticReport.docx")
        if os.path.exists(os.path.dirname(sys.argv[0]) + r"\\Report_Word\\" + testName + "Post-StaticReport.docx"):
            files.append(testName + "Post-StaticReport.docx")
        files.append(testName + "RawData.docx")
        self.combine_word_documents(files, testName)
        self.close()

