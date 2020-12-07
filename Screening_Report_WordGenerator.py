import pandas as pd
import numpy as np
import docx
import os
import datetime
from docx import Document

pd.options.display.max_columns = 999
pd.options.display.max_rows = 999
pd.set_option("display.precision", 6)


def combine_word_documents(files):
    merged_document = Document()

    for index, file in enumerate(files):
        sub_doc = Document(os.getcwd() + r"\\Report_Word\\" + file)

        # Don't add a page break if you've reached the last file.
        if index < len(files)-1:
            sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
        print(os.getcwd() + r"\\Report_Word\\" + 'merged' + files[0])
    merged_document.save(os.getcwd() + r"\\Report_Word\\" + 'merged' + files[0])

def screening_Report_wordGenerator(x):
    # import the template
    doc = docx.Document(os.getcwd() + r"\\Report_Word\\Doc Template\\demo.docx")
    # get the template info
    path_template = os.getcwd() + r"\\Screening_Template\\" + x
    data_file1 = pd.read_csv(path_template, sep="\t")
    columns_1 = data_file1.loc[0].fillna("")
    # get the testing data
    path_datafile = os.getcwd() + r"\\Screening_Data\\" + x
    data_file = pd.read_csv(path_datafile, sep="\t")

    '''
    i = 0
    for text in doc.paragraphs:
        run_i = 0
        for runs in text.runs:
            print(runs.text)
            print("This is " + str(i) + " lin and " + str(run_i) + " runs")
            run_i = run_i + 1
        i = i + 1
    '''
    # Test Number
    doc.paragraphs[1].runs[1].text = str(x[:-4])
    # Report Data
    doc.paragraphs[1].runs[6].text = str(datetime.date.today())
    # Battery Name
    doc.paragraphs[2].runs[1].text = columns_1["Cell Name"]
    # Chemistry Name
    doc.paragraphs[3].runs[1].text = columns_1["Chemistry"]
    # request Number
    doc.paragraphs[5].runs[5].text = str(columns_1["Request Number"])
    # Test Number
    doc.paragraphs[6].runs[6].text = str(columns_1["Task Number"])
    # Lot Number
    doc.paragraphs[7].runs[6].text = str(columns_1["Lot No"])
    # Form Factor
    doc.paragraphs[8].runs[7].text = columns_1["Form Factor"]
    # Test Purpose
    doc.paragraphs[9].runs[6].text = columns_1["Test Purpose"]
    # Project Engineer
    doc.paragraphs[10].runs[5].text = columns_1["Tech POC."]
    # find the right crtiria
    if columns_1["Tabbed?"] == "Tabbed":
        doc.paragraphs[11].runs[5].text = "      PRE TAB"
        doc.paragraphs[11].runs[10].text = "     POST TAB"
        # pretabe-OCV volt for tabbed
        doc.paragraphs[12].runs[6].text = str(columns_1["Pre-Tab OCV"])
        # post Tab OCV volt
        doc.paragraphs[12].runs[11].text = str(columns_1["Post-Tab OCV"])
        # pre tab CCV
        doc.paragraphs[13].runs[6].text = "N/A"
        # post CCV
        doc.paragraphs[13].runs[11].text = str(columns_1["Post-Tab CCV"])
        # Pre tab Drain time
        doc.paragraphs[14].runs[6].text = "N/A"
        # Post Tab Drain Time
        doc.paragraphs[14].runs[11].text = str(columns_1["Profile One Timer"])
        # load type
        if columns_1["Profile One Type"] == "Constant Current":
            doc.paragraphs[15].runs[0].text = "Current(mA)"
        elif columns_1["Profile One Type"] == "Constant Resistor":
            doc.paragraphs[15].runs[0].text = "Resistor (Ω): "
        doc.paragraphs[15].runs[7].text = "N/A"
        doc.paragraphs[15].runs[12].text = str(columns_1["Profile One Values"])
        # OCV Tolerance
        doc.paragraphs[16].runs[3].text = str(columns_1["OCV Tab Tolerance"])
    elif columns_1["Section No."] == 2:
        doc.paragraphs[11].runs[5].text = " Section One"
        doc.paragraphs[11].runs[10].text = " Section Two"
        # pre tab-OCV volt for tabbed
        doc.paragraphs[12].runs[6].text = str(columns_1["Profile One OCV Min"])
        # post Tab OCV volt
        doc.paragraphs[12].runs[11].text = str(columns_1["Profile One OCV Min"])
        # pre tab CCV
        doc.paragraphs[13].runs[6].text = str(columns_1["Profile One CCV Min"])
        # post CCV
        doc.paragraphs[13].runs[11].text = str(columns_1["Profile One CCV Min"])
        # Pre tab Drain time
        doc.paragraphs[14].runs[6].text = str(columns_1["Profile One Timer"])
        # Post Tab Drain Time
        doc.paragraphs[14].runs[11].text = str(columns_1["Profile One Timer"])
        # load type
        if columns_1["Profile One Type"] == "Constant Current":
            doc.paragraphs[15].runs[0].text = "Current(mA)"
        elif columns_1["Profile One Type"] == "Constant Resistor":
            doc.paragraphs[15].runs[0].text = "Resistor (Ω): "
        doc.paragraphs[15].runs[7].text = str(columns_1["Profile One Values"])
        doc.paragraphs[15].runs[12].text = str(columns_1["Profile One Values"])
        # make OCV Tolerance empty
        doc.paragraphs[16].runs[0].text = ""
        doc.paragraphs[16].runs[1].text = ""
        doc.paragraphs[16].runs[2].text = ""
        doc.paragraphs[16].runs[3].text = ""
        doc.paragraphs[16].runs[4].text = ""
        doc.paragraphs[16].runs[5].text = ""
    elif columns_1["Profile No."] == 2:
        doc.paragraphs[11].runs[5].text = "Profile One"
        doc.paragraphs[11].runs[10].text = "Profile Two"
        # pre tab-OCV volt for tabbed
        doc.paragraphs[12].runs[6].text = str(columns_1["Profile One OCV Min"])
        # post Tab OCV volt
        doc.paragraphs[12].runs[11].text = str(columns_1["Profile Two OCV Min"])
        # pre tab CCV
        doc.paragraphs[13].runs[6].text = str(columns_1["Profile One CCV Min"])
        # post CCV
        doc.paragraphs[13].runs[11].text = str(columns_1["Profile Two CCV Min"])
        # Pre tab Drain time
        doc.paragraphs[14].runs[6].text = str(columns_1["Profile One Timer"])
        # Post Tab Drain Time
        doc.paragraphs[14].runs[11].text = str(columns_1["Profile Two Timer"])
        # load type
        if columns_1["Profile One Type"] == "Constant Current":
            doc.paragraphs[15].runs[0].text = "Current(mA)"
        elif columns_1["Profile One Type"] == "Constant Resistor":
            doc.paragraphs[15].runs[0].text = "Resistor (Ω): "
        doc.paragraphs[15].runs[7].text = str(columns_1["Profile One Values"])
        doc.paragraphs[15].runs[12].text = str(columns_1["Profile Two Value"])
        # make OCV Tolerance empty
        doc.paragraphs[16].runs[0].text = ""
        doc.paragraphs[16].runs[1].text = ""
        doc.paragraphs[16].runs[2].text = ""
        doc.paragraphs[16].runs[3].text = ""
        doc.paragraphs[16].runs[4].text = ""
        doc.paragraphs[16].runs[5].text = ""
    else:
        doc.paragraphs[11].runs[5].text = "Screening"
        doc.paragraphs[11].runs[10].text = ""
        doc.paragraphs[11].runs[12].text = ""
        # pretabe-OCV volt for tabbed
        doc.paragraphs[12].runs[6].text = str(columns_1["Profile One OCV Min"])
        # post Tab OCV volt
        doc.paragraphs[12].runs[11].text = ""
        # pre tab CCV
        doc.paragraphs[13].runs[6].text = str(columns_1["Profile One CCV Min"])
        # post CCV
        doc.paragraphs[13].runs[11].text = ""
        # Pre tab Drain time
        doc.paragraphs[14].runs[6].text = str(columns_1["Profile One Timer"])
        # Post Tab Drain Time
        doc.paragraphs[14].runs[11].text = ""
        # load type
        if columns_1["Profile One Type"] == "Constant Current":
            doc.paragraphs[15].runs[0].text = "Current(mA)"
        elif columns_1["Profile One Type"] == "Constant Resistor":
            doc.paragraphs[15].runs[0].text = "Resistor (Ω): "
        doc.paragraphs[15].runs[7].text = str(columns_1["Profile One Values"])
        doc.paragraphs[15].runs[12].text = ""
        # make OCV Tolerance empty
        doc.paragraphs[16].runs[0].text = ""
        doc.paragraphs[16].runs[1].text = ""
        doc.paragraphs[16].runs[2].text = ""
        doc.paragraphs[16].runs[3].text = ""
        doc.paragraphs[16].runs[4].text = ""
        doc.paragraphs[16].runs[5].text = ""
    # Battery Size
    doc.paragraphs[18].runs[8].text = str(columns_1["Dimension (mm)"])
    # Discharge Temperautre
    doc.paragraphs[19].runs[4].text = str(columns_1["Screening Temp(°C)"])
    # Number of batteries Tested
    doc.paragraphs[21].runs[4].text = str(len(data_file["Barcode"]))
    # Number of the batteries failled
    if data_file["Post-OCV"].isnull().all():
        doc.paragraphs[22].runs[6].text = str(len(data_file[data_file["Pre-screen pass"] == "Fail"]))
    else:
        # ! case not senstive
        pre_ScreenPass = data_file[data_file["Pre-screen pass"].dropna() == "Pass"]["Post-screen pass"]
        doc.paragraphs[22].runs[6].text = str(len(data_file["Barcode"]) - len(pre_ScreenPass[pre_ScreenPass == "Pass"]))

    # Date of Manufacture:
    doc.paragraphs[24].runs[5].text = str(columns_1["Mfg Date"])
    # Date of recepit:
    doc.paragraphs[25].runs[5].text = str(columns_1["Date Received"])
    # Date started:
    doc.paragraphs[26].runs[6].text = str(columns_1["Begin Date"])
    # Date finished
    doc.paragraphs[27].runs[5].text = str(columns_1["Finished Date"])
    # save the doc
    report_name = x[:-3]+"docx"
    doc.save(os.getcwd() + r"\\Report_Word\\" + report_name)

    # generate the Note doc
    doc1 = docx.Document(os.getcwd() + r"\\Report_Word\\Doc Template\\Note-template.docx")
    # Test number
    doc1.paragraphs[0].runs[1].text = str(x[:-4])
    # add note
    doc1.paragraphs[1].runs[0].text = str(columns_1["Note"])
    note_name = "Note-" + x[:-3] + "docx"
    doc1.save(os.getcwd() + r"\\Report_Word\\" + note_name)
    files = [report_name, note_name]
    combine_word_documents(files)


# for debugging propose
if __name__ == '__main__':
    screening_Report_wordGenerator("14575A00.txt")