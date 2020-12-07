import os
import pyodbc
import datetime
import pandas as pd
import math
import numpy as np
import sys
from PyQt5.QtWidgets import QApplication, QDialog, QFileDialog, QMessageBox, QTableWidgetItem
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.uic import loadUi
from Screening_System_PyQt5 import client
from PyQt5.QtCore import Qt
from PyQt5.uic import loadUi
from scipy import stats
from docx import Document
import docx


pd.options.display.max_columns = 999
pd.options.display.max_rows = 999
pd.set_option("display.precision", 6)


class Report_StatisticsTable_WithFunctions(QDialog):

    def __init__(self):
        super(Report_StatisticsTable_WithFunctions, self).__init__()
        loadUi("RGui_Report_StatisticsTable.ui", self)

    def count(self, list1, l, r):
        return len(list(x for x in list1 if l <= x <= r))

    def samplePassingTabTolerance(self, data_file, column_1, y):
        data_file2 = data_file[data_file["Pre-OCV"] >= column_1["Pre-Tab OCV"]]
        # pw is how many cell that has PreOCV and PostOCV deference bigger than OCV Tab Tolerance
        allData = round(data_file2["Post-OCV"].dropna(), 3)
        passingData = allData[allData >= column_1["Post-Tab OCV"]]
        ocvpassingdata = data_file2[data_file2["Post-OCV"] >= column_1["Post-Tab OCV"]]["Pre-OCV"]
        print("Post-Tab OCV " + str(column_1["Post-Tab OCV"]))
        pdata = {"OCV": ocvpassingdata, "POCV": passingData}
        df = pd.DataFrame(pdata)
        df = df[round(abs(df["OCV"] - df["POCV"]), 3) <= column_1["OCV Tab Tolerance"]]
        if y == "Pre-OCV":
            df = df["OCV"]
        elif y == "Post-OCV":
            df = df["POCV"]
        elif y == "Post-CCV":
            df = data_file["Post-CCV"].dropna()
            df = df[df >= column_1["Post-Tab CCV"]]

        binList = []
        countList = []
        df = round(df, 3)
        MaxtoMin = round(float(max(df)) - float(min(df)), 3)
        HowManyBin = int(round(math.sqrt(df.size))) + 2
        histgram_BinWith = round((MaxtoMin / HowManyBin), 3) + 0.001
        # if the sample size is too small, this will make sure the binWith is aleast 0.001
        if histgram_BinWith < 0.001:
            histgram_BinWith = 0.001
        lowest_point = min(df) - histgram_BinWith
        eachBin = max(df)
        while eachBin >= lowest_point:
            lowBin = round(eachBin - round(histgram_BinWith, 3), 3)
            highBin = round(eachBin, 3)
            if round(eachBin - round(histgram_BinWith, 3), 3) <= lowest_point:
                binList.append(str("%.3f" % round(highBin, 3)) + "-" + str("%.3f" % round(lowBin, 3)))
                countList.append(self.count(df, lowBin, highBin))
            else:
                binList.append(str("%.3f" % round(highBin, 3)) + "-" + str("%.3f" % round(lowBin, 3)))
                countList.append(self.count(df, lowBin + 0.001, highBin))
            eachBin = round(eachBin - round(histgram_BinWith, 3), 3)

        return binList, countList

    def screening_Report_partone(self, doc, x, column_1, table_one, label_4):
        '''
        i = 0
        # the paragraphs
        print("Checking321")
        for text in doc.paragraphs:
            run_i = 0
            for runs in text.runs:
                print(runs.text)
                print("This is " + str(i) + " lin and " + str(run_i) + " runs")
                run_i = run_i + 1
            i = i + 1
        '''
        # se the testNumber
        doc.paragraphs[1].runs[1].text = str(x)
        # set the date
        doc.paragraphs[1].runs[9].text = str(datetime.date.today())
        # set name
        doc.paragraphs[2].runs[2].text = str(column_1["Cell Name"])
        # set chemistry
        SUB = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
        doc.paragraphs[3].runs[2].text = str(column_1["Chemistry"].translate(SUB))
        # set Name
        doc.paragraphs[4].runs[0].text = str(table_one.horizontalHeaderItem(0).text()[: -4])
        # pre-tab report for Pre-Tab OCv
        if table_one.horizontalHeaderItem(0).text() == "Pre-Tab OCV":
            # set ocv max
            doc.paragraphs[6].runs[4].text = str(table_one.item(1, 0).text())
            # set ocv min
            doc.paragraphs[7].runs[4].text = str(table_one.item(2, 0).text())
            # set ocv median
            doc.paragraphs[8].runs[4].text = str(table_one.item(3, 0).text())
            # set ocv mean
            doc.paragraphs[9].runs[4].text = str(table_one.item(4, 0).text())
            # set ocv stander deviation
            doc.paragraphs[10].runs[4].text = str(table_one.item(5, 0).text())
            # set ocv total samples
            doc.paragraphs[11].runs[4].text = str(table_one.item(9, 0).text())
            # set ocv total pass_criteriion
            doc.paragraphs[12].runs[3].text = str(
                int(float(table_one.item(9, 0).text()) - float(table_one.item(10, 0).text())))
            # set ocv total failling criteria
            doc.paragraphs[13].runs[3].text = str(table_one.item(10, 0).text())
            # set ocv 95% confidence Interval
            doc.paragraphs[14].runs[3].text = str(table_one.item(6, 0).text()) + "/" + str(table_one.item(7, 0).text())
            # set total sample outside the range
            doc.paragraphs[15].runs[3].text = str(table_one.item(8, 0).text())
            # set OCV critria
            doc.paragraphs[12].runs[6].text = str(table_one.item(0, 0).text())

        if table_one.horizontalHeaderItem(0).text() == "Profile One OCV" or \
                table_one.horizontalHeaderItem(0).text() == "Section One OCV" or \
                table_one.horizontalHeaderItem(0).text() == "OCV":
            # set ocv max
            doc.paragraphs[6].runs[5].text = str(table_one.item(1, 0).text())
            # set ocv min
            doc.paragraphs[7].runs[6].text = str(table_one.item(2, 0).text())
            # set ocv median
            doc.paragraphs[8].runs[5].text = str(table_one.item(3, 0).text())
            # set ocv mean
            doc.paragraphs[9].runs[5].text = str(table_one.item(4, 0).text())
            # set ocv stander deviation
            doc.paragraphs[10].runs[5].text = str(table_one.item(5, 0).text())
            # set ocv total samples
            doc.paragraphs[11].runs[5].text = str(table_one.item(9, 0).text())
            # set ocv total pass_criteriion
            doc.paragraphs[12].runs[4].text = str(
                int(float(table_one.item(9, 0).text()) - float(table_one.item(10, 0).text())))
            # set ocv total failling criteria
            doc.paragraphs[13].runs[4].text = str(table_one.item(10, 0).text())
            # set ocv 95% confidence Interval
            doc.paragraphs[14].runs[5].text = str(table_one.item(6, 0).text()) + "/" + str(table_one.item(7, 0).text())
            # set total sample outside the range
            doc.paragraphs[15].runs[4].text = str(table_one.item(8, 0).text())
            # set OCV critria
            doc.paragraphs[16].runs[0].text = "OCV >or= " + str(table_one.item(0, 0).text())

        if table_one.horizontalHeaderItem(1).text() == "Profile One CCV" or \
                table_one.horizontalHeaderItem(1).text() == "Section One CCV" or \
                table_one.horizontalHeaderItem(1).text() == "CCV":
            doc.paragraphs[4].runs[0].text = str(table_one.horizontalHeaderItem(2).text()[: -4])
            # CCV max
            doc.paragraphs[6].runs[9].text = str(table_one.item(1, 1).text())
            # CCV min
            doc.paragraphs[7].runs[10].text = str(table_one.item(2, 1).text())
            # Median
            doc.paragraphs[8].runs[9].text = str(table_one.item(3, 1).text())
            # Means
            doc.paragraphs[9].runs[9].text = str(table_one.item(4, 1).text())
            # stander deviaton
            doc.paragraphs[10].runs[9].text = str(table_one.item(5, 1).text())
            # total samples
            doc.paragraphs[11].runs[9].text = str(table_one.item(9, 1).text())
            # total pass criterion
            doc.paragraphs[12].runs[8].text = str(int(float(table_one.item(9, 1).text()) - float(table_one.item(10, 1).text())))
            # total failing criterion
            doc.paragraphs[13].runs[8].text = str(table_one.item(10, 1).text())
            # set ocv 95% confidence Interval
            doc.paragraphs[14].runs[8].text = str(table_one.item(6, 1).text()) + "/" + str(table_one.item(7, 1).text())
            # set total sample outside the range
            doc.paragraphs[15].runs[8].text = str(table_one.item(8, 1).text())
            # set OCV critria
            doc.paragraphs[16].runs[0].text = "OCV >or= " + str(table_one.item(0, 0).text())
            # set CCV criteria
            doc.paragraphs[17].runs[0].text = "               CCV >or= " + str(table_one.item(0, 1).text()) + " @ " + str(label_4.text())

    def screening_Report_partTwo(self, doc, x, column_1, table_one, label_5):

        i = 0
        # the paragraphs
        print("Checking321")
        for text in doc.paragraphs:
            run_i = 0
            for runs in text.runs:
                print(runs.text)
                print("This is " + str(i) + " lin and " + str(run_i) + " runs")
                run_i = run_i + 1
            i = i + 1

        # pre-tab report for Pre-Tab OCv
        if table_one.horizontalHeaderItem(0).text() == "Pre-Tab OCV":
            # se the testNumber
            doc.paragraphs[1].runs[1].text = str(x)
            # set the date
            doc.paragraphs[1].runs[9].text = str(datetime.date.today())
            # set name
            doc.paragraphs[2].runs[2].text = str(column_1["Cell Name"])
            # set chemistry
            SUB = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
            doc.paragraphs[3].runs[2].text = str(column_1["Chemistry"].translate(SUB))
            # set Name
            doc.paragraphs[4].runs[0].text = str(table_one.horizontalHeaderItem(2).text()[: -4])
            # post-OCV report
            # set ocv max
            doc.paragraphs[6].runs[3].text = str(table_one.item(1, 2).text())
            # set ocv min
            doc.paragraphs[7].runs[3].text = str(table_one.item(2, 2).text())
            # set ocv median
            doc.paragraphs[8].runs[4].text = str(table_one.item(3, 2).text())
            # set ocv mean
            doc.paragraphs[9].runs[3].text = str(table_one.item(4, 2).text())
            # set ocv stander deviation
            doc.paragraphs[10].runs[3].text = str(table_one.item(5, 2).text())
            # set ocv total samples
            doc.paragraphs[11].runs[3].text = str(table_one.item(9, 2).text())
            # set ocv total pass_criteriion
            doc.paragraphs[12].runs[2].text = str(
                int(float(table_one.item(9, 2).text()) - float(table_one.item(10, 2).text())))
            # set ocv total failling criteria
            doc.paragraphs[13].runs[2].text = str(table_one.item(10, 2).text())
            # set ocv 95% confidence Interval
            doc.paragraphs[14].runs[3].text = str(table_one.item(6, 2).text()) + "/" + str(table_one.item(7, 2).text())
            # set total sample outside the range
            doc.paragraphs[15].runs[5].text = str(table_one.item(8, 2).text())
            # set OCV critria
            doc.paragraphs[11].runs[11].text = str(table_one.item(0, 2).text())

            # post_CCV report
            # CCV max
            doc.paragraphs[6].runs[7].text = str(table_one.item(1, 3).text()) + "   "
            # CCV min
            doc.paragraphs[7].runs[6].text = "    " + str(table_one.item(2, 3).text()) + "  "
            # Median
            doc.paragraphs[8].runs[7].text = "    " + str(table_one.item(3, 3).text()) + "  "
            # Means
            doc.paragraphs[9].runs[7].text = str(table_one.item(4, 3).text())
            doc.paragraphs[9].runs[8].text = "  "
            # stander deviaton
            doc.paragraphs[10].runs[6].text = "    " + str(table_one.item(5, 3).text())
            # total samples
            doc.paragraphs[11].runs[7].text = str(table_one.item(9, 3).text()) + "      "
            # total pass criterion
            doc.paragraphs[12].runs[5].text = str(
                int(float(table_one.item(9, 3).text()) - float(table_one.item(10, 3).text()))) + "      "
            # total failing criterion
            doc.paragraphs[13].runs[6].text = str(table_one.item(10, 3).text()) + "         "
            # set ocv 95% confidence Interval
            doc.paragraphs[14].runs[6].text = "    " + str(table_one.item(6, 3).text()) + "/" + str(table_one.item(7, 3).text())
            doc.paragraphs[14].runs[7].text = ""
            doc.paragraphs[14].runs[8].text = ""
            doc.paragraphs[14].runs[9].text = ""
            # set total sample outside the range Peter Change
            doc.paragraphs[15].runs[9].text = str(table_one.item(8, 3).text())
            # set CCV criteria
            doc.paragraphs[12].runs[10].text = str(table_one.item(0, 3).text())
            doc.paragraphs[12].runs[13].text = str(label_5.text())
            # set tab tolerance
            doc.paragraphs[14].runs[13].text = str(column_1["OCV Tab Tolerance"])
        else:
            # se the testNumber
            doc.paragraphs[1].runs[1].text = str(x)
            # set the date
            doc.paragraphs[1].runs[9].text = str(datetime.date.today())
            # set name
            doc.paragraphs[2].runs[2].text = str(column_1["Cell Name"])
            # set
            SUB = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
            doc.paragraphs[3].runs[2].text = str(column_1["Chemistry"].translate(SUB))
            # set Name
            doc.paragraphs[4].runs[0].text = str(table_one.horizontalHeaderItem(2).text()[: -4])
            # post-OCV report
            # set ocv max
            doc.paragraphs[6].runs[5].text = str(table_one.item(1, 2).text())
            # set ocv min
            doc.paragraphs[7].runs[6].text = str(table_one.item(2, 2).text())
            # set ocv median
            doc.paragraphs[8].runs[5].text = str(table_one.item(3, 2).text())
            # set ocv mean
            doc.paragraphs[9].runs[5].text = str(table_one.item(4, 2).text())
            # set ocv stander deviation
            doc.paragraphs[10].runs[5].text = str(table_one.item(5, 2).text())
            # set ocv total samples
            doc.paragraphs[11].runs[5].text = str(table_one.item(9, 2).text())
            # set ocv total pass_criteriion
            doc.paragraphs[12].runs[4].text = str(
                int(float(table_one.item(9, 2).text()) - float(table_one.item(10, 2).text())))
            # set ocv total failling criteria
            doc.paragraphs[13].runs[4].text = str(table_one.item(10, 2).text())
            # set ocv 95% confidence Interval
            doc.paragraphs[14].runs[5].text = str(table_one.item(6, 2).text()) + "/" + str(table_one.item(7, 2).text())
            # set total sample outside the range
            doc.paragraphs[15].runs[4].text = str(table_one.item(8, 2).text())
            # set OCV critria
            doc.paragraphs[16].runs[0].text = "OCV >or= " + str(table_one.item(0, 2).text())
            # post_CCV report
            # CCV max
            doc.paragraphs[6].runs[9].text = str(table_one.item(1, 3).text())
            # CCV min
            doc.paragraphs[7].runs[10].text = str(table_one.item(2, 3).text())
            # Median
            doc.paragraphs[8].runs[9].text = str(table_one.item(3, 3).text())
            # Means
            doc.paragraphs[9].runs[9].text = str(table_one.item(4, 3).text())
            # stander deviaton
            doc.paragraphs[10].runs[9].text = str(table_one.item(5, 3).text())
            # total samples
            doc.paragraphs[11].runs[9].text = str(table_one.item(9, 3).text())
            # total pass criterion
            doc.paragraphs[12].runs[8].text = str(int(float(table_one.item(9, 3).text()) - float(table_one.item(10, 3).text())))
            # total failing criterion
            doc.paragraphs[13].runs[8].text = str(table_one.item(10, 3).text())
            # set ocv 95% confidence Interval
            doc.paragraphs[14].runs[8].text = str(table_one.item(6, 3).text()) + "/" + str(table_one.item(7, 3).text())
            # set total sample outside the range Peter Change
            doc.paragraphs[15].runs[8].text = str(table_one.item(8, 3).text())
            # set CCV criteria
            doc.paragraphs[17].runs[0].text = "               CCV >or= " + str(table_one.item(0, 3).text()) + " @ " + str(label_5.text())

    def screening_Report_wordGenerator(self, x, column_1, table_one, table_four, table_two, table_five, table_three, label_4, label_5):
        print(column_1)
        # to use Pre-Tab template
        if table_one.horizontalHeaderItem(0).text() == "Pre-Tab OCV":
            # report for the pre-tab
            doc = docx.Document(os.getcwd() + r"\\Report_Word\\Doc Template\\Static-template-PreTab-New.docx")
            self.screening_Report_partone(doc, x, column_1, table_one, label_4)
            PreTab_table = doc.tables[0]
            # load the data for passing samples
            path_datafile = os.getcwd() + r"\\Screening_Data\\" + x + ".txt"
            data_file = pd.read_csv(path_datafile, sep="\t")
            # finding the pass data file
            data_file2 = data_file[data_file["Pre-OCV"] >= column_1["Pre-Tab OCV"]]
            # pw is how many cell that has PreOCV and PostOCV deference bigger than OCV Tab Tolerance
            allData = round(data_file2["Post-OCV"].dropna(), 3)
            passingData = allData[allData >= column_1["Post-Tab OCV"]]
            ocvpassingdata = data_file2[data_file2["Post-OCV"] >= column_1["Post-Tab OCV"]]["Pre-OCV"]
            print("Post-Tab OCV " + str(column_1["Post-Tab OCV"]))
            pdata = {"OCV": ocvpassingdata, "POCV": passingData}
            df = pd.DataFrame(pdata)
            df = df[round(abs(df["OCV"] - df["POCV"]), 3) <= column_1["OCV Tab Tolerance"]]

            if not data_file["Pre-OCV"].isnull().all():
                df3 = df["OCV"]
                # Pre-tab OCV passing sample OCV maximun
                doc.paragraphs[6].runs[10].text = str("%.3f" % max(df3))
                # pre-tab OCV passing sample OCV min
                doc.paragraphs[7].runs[12].text = str("%.3f" % min(df3))
                # Pre_tab OCV Median
                doc.paragraphs[8].runs[10].text = str("%.3f" % round(df3.median(), 3))
                # pre_tabe OCV Mean
                doc.paragraphs[9].runs[12].text = str("%.3f" % round(df3.mean(), 3))
                # pre_tab OCV Stander Deivation
                doc.paragraphs[10].runs[11].text = str(round(df3.std(), 5))
                # set the Tab Tolerance
                doc.paragraphs[14].runs[5].text = str(column_1["OCV Tab Tolerance"])

            sampleTabTolerance = self.samplePassingTabTolerance(data_file, column_1, "Pre-OCV")

            for i in range(table_four.rowCount()):
                # check if the item is none type
                if type(table_four.item(i, 0)) is type(table_four.item(0, 0)):
                    cell_info = PreTab_table.add_row()
                    cell_info.cells[0].text = str(table_four.item(i, 0).text()) + "-" + str(table_four.item(i, 1).text())
                    cell_info.cells[1].text = str(table_four.item(i, 2).text())
                    if i < len(sampleTabTolerance[0]):
                        cell_info.cells[2].text = str(sampleTabTolerance[0][i])
                        cell_info.cells[3].text = str(sampleTabTolerance[1][i])

            doc.save(os.getcwd() + r"\\Report_Word\\" + x + 'Pre-StaticReport.docx')

            # report for the post
            doc1 = docx.Document(os.getcwd() + r"\\Report_Word\\Doc Template\\Static-template-PostTab-New.docx")
            self.screening_Report_partTwo(doc1, x, column_1, table_one, label_5)

            # load the data for passing samples
            path_datafile = os.getcwd() + r"\\Screening_Data\\" + x + ".txt"
            data_file = pd.read_csv(path_datafile, sep="\t")
            if not data_file["Post-OCV"].isnull().all():
                df1 = df["POCV"]
                # Post-tab OCV passing sample OCV maximun
                doc1.paragraphs[6].runs[12].text = str("%.3f" % max(df1))
                doc1.paragraphs[6].runs[13].text = ""
                # post-tab OCV passing sample OCV min
                doc1.paragraphs[7].runs[13].text = str("%.3f" % min(df1))
                # Post_tab OCV Median
                doc1.paragraphs[8].runs[14].text = str("%.3f" % round(df1.median(), 3))
                # post_tabe OCV Mean
                doc1.paragraphs[9].runs[15].text = str("%.3f" % round(df1.mean(), 3))
                # post_tab OCV Stander Deivation
                doc1.paragraphs[10].runs[12].text = str(round(df1.std(), 5))
            if not data_file["Post-CCV"].isnull().all():
                df4 = data_file["Post-CCV"].dropna()
                df4 = df4[df4 >= column_1["Post-Tab CCV"]]
                # Post-tab OCV passing sample OCV maximun
                doc1.paragraphs[6].runs[16].text = str("%.3f" % max(df4))
                # post-tab OCV passing sample OCV min
                doc1.paragraphs[7].runs[16].text = str("%.3f" % min(df4))
                # Post_tab OCV Median
                doc1.paragraphs[8].runs[17].text = str("%.3f" % round(df4.median(), 3))
                # post_tabe OCV Mean
                doc1.paragraphs[9].runs[18].text = str("%.3f" % round(df4.mean(), 3))
                # post_tab OCV Stander Deivation
                doc1.paragraphs[10].runs[15].text = str(round(df4.std(), 5))

            sampleTabTolerance_post_OCV = self.samplePassingTabTolerance(data_file, column_1, "Post-OCV")
            sampleTabTolerance_post_CCV = self.samplePassingTabTolerance(data_file, column_1, "Post-CCV")
            PostTab_table = doc1.tables[0]
            if table_five.rowCount() > table_three.rowCount():
                HowmanyTablerows = table_five.rowCount()
            else:
                HowmanyTablerows = table_three.rowCount()
            for i in range(HowmanyTablerows):
                # check if the item is none type
                cell_info = PostTab_table.add_row()
                if type(table_five.item(i, 0)) is type(table_five.item(0, 0)):
                    cell_info.cells[0].text = str(table_five.item(i, 0).text()) + "-" + str(
                        table_five.item(i, 1).text())
                    cell_info.cells[1].text = str(table_five.item(i, 2).text())
                else:
                    cell_info.cells[0].text = ""
                    cell_info.cells[1].text = ""
                if type(table_three.item(i, 0)) is type(table_three.item(0, 0)):
                    cell_info.cells[2].text = str(table_three.item(i, 0).text()) + "-" + str(
                        table_three.item(i, 1).text())
                    cell_info.cells[3].text = str(table_three.item(i, 2).text())
                else:
                    cell_info.cells[2].text = ""
                    cell_info.cells[3].text = ""
                if i < len(sampleTabTolerance_post_OCV[0]):
                    cell_info.cells[4].text = str(sampleTabTolerance_post_OCV[0][i])
                    cell_info.cells[5].text = str(sampleTabTolerance_post_OCV[1][i])
                if i < len(sampleTabTolerance_post_CCV[0]):
                    cell_info.cells[6].text = str(sampleTabTolerance_post_CCV[0][i])
                    cell_info.cells[7].text = str(sampleTabTolerance_post_CCV[1][i])

            doc1.save(os.getcwd() + r"\\Report_Word\\" + x + 'Post-StaticReport.docx')

        elif table_one.horizontalHeaderItem(0).text() == "OCV":
            doc = docx.Document(os.getcwd() + r"\\Report_Word\\Doc Template\\Static-template.docx")
            self.screening_Report_partone(doc, x, column_1, table_one, label_4)

            PostTab_table = doc.tables[0]

            if table_four.rowCount() > table_two.rowCount():
                HowmanyTablerows = table_four.rowCount()
            else:
                HowmanyTablerows = table_two.rowCount()
            for i in range(HowmanyTablerows):
                # check if the item is none type
                cell_info = PostTab_table.add_row()
                if type(table_four.item(i, 0)) is type(table_four.item(0, 0)):
                    cell_info.cells[0].text = str(table_four.item(i, 0).text()) + "-" + str(
                        table_four.item(i, 1).text())
                    cell_info.cells[1].text = str(table_four.item(i, 2).text())
                else:
                    cell_info.cells[0].text = ""
                    cell_info.cells[1].text = ""
                if type(table_two.item(i, 0)) is type(table_two.item(0, 0)):
                    cell_info.cells[2].text = str(table_two.item(i, 0).text()) + "-" + str(
                        table_two.item(i, 1).text())
                    cell_info.cells[3].text = str(table_two.item(i, 2).text())
                else:
                    cell_info.cells[2].text = ""
                    cell_info.cells[3].text = ""
            doc.save(os.getcwd() + r"\\Report_Word\\" + x + 'Pre-StaticReport.docx')
        else:
            # pre_report
            doc = docx.Document(os.getcwd() + r"\\Report_Word\\Doc Template\\Static-template.docx")
            self.screening_Report_partone(doc, x, column_1, table_one, label_4)

            PostTab_table = doc.tables[0]

            if table_four.rowCount() > table_two.rowCount():
                HowmanyTablerows = table_four.rowCount()
            else:
                HowmanyTablerows = table_two.rowCount()
            for i in range(HowmanyTablerows):
                # check if the item is none type
                cell_info = PostTab_table.add_row()
                if type(table_four.item(i, 0)) is type(table_four.item(0, 0)):
                    cell_info.cells[0].text = str(table_four.item(i, 0).text()) + "-" + str(
                        table_four.item(i, 1).text())
                    cell_info.cells[1].text = str(table_four.item(i, 2).text())
                else:
                    cell_info.cells[0].text = ""
                    cell_info.cells[1].text = ""
                if type(table_two.item(i, 0)) is type(table_two.item(0, 0)):
                    cell_info.cells[2].text = str(table_two.item(i, 0).text()) + "-" + str(
                        table_two.item(i, 1).text())
                    cell_info.cells[3].text = str(table_two.item(i, 2).text())
                else:
                    cell_info.cells[2].text = ""
                    cell_info.cells[3].text = ""
            doc.save(os.getcwd() + r"\\Report_Word\\" + x + 'Pre-StaticReport.docx')

            # report for the post
            doc1 = docx.Document(os.getcwd() + r"\\Report_Word\\Doc Template\\Static-template.docx")
            self.screening_Report_partTwo(doc1, x, column_1, table_one, label_5)
            PostTab_table = doc1.tables[0]

            if table_five.rowCount() > table_three.rowCount():
                HowmanyTablerows = table_five.rowCount()
            else:
                HowmanyTablerows = table_three.rowCount()
            for i in range(HowmanyTablerows):
                # check if the item is none type
                cell_info = PostTab_table.add_row()
                if type(table_five.item(i, 0)) is type(table_five.item(0, 0)):
                    cell_info.cells[0].text = str(table_five.item(i, 0).text()) + "-" + str(
                        table_five.item(i, 1).text())
                    cell_info.cells[1].text = str(table_five.item(i, 2).text())
                else:
                    cell_info.cells[0].text = ""
                    cell_info.cells[1].text = ""
                if type(table_three.item(i, 0)) is type(table_three.item(0, 0)):
                    cell_info.cells[2].text = str(table_three.item(i, 0).text()) + "-" + str(
                        table_three.item(i, 1).text())
                    cell_info.cells[3].text = str(table_three.item(i, 2).text())
                else:
                    cell_info.cells[2].text = ""
                    cell_info.cells[3].text = ""
            doc1.save(os.getcwd() + r"\\Report_Word\\" + x + 'Post-StaticReport.docx')

    def findSampleOutlier(self, data):
        iqr = stats.iqr(data)
        q1, q3 = np.percentile(data, [25, 75])
        min = round(q1 - (1.5 * iqr), 3)
        max = round(q3 + (1.5 * iqr), 3)
        outlier = len(data) - len(data[data.between(min, max)])
        return outlier

    def mean_confidence_interval(self, data, confidence=0.95):
        a = 1.0 * np.array(data)
        n = len(a)
        m, se = np.mean(a), stats.sem(a)
        h = se * stats.t.ppf((1 + confidence) / 2., n - 1)
        # m is mean, m-h is min, m+h is max
        return m, m - h, m + h

    def getTestNumber2(self, x):
        path_datafile = os.getcwd() + r"\\Screening_Data\\" + x
        print(path_datafile)
        data_file = pd.read_csv(path_datafile, sep="\t")
        #Load the template info for report
        path_template = os.getcwd() + r"\\Screening_Template\\" + x
        print(path_template)
        data_file1 = pd.read_csv(path_template, sep="\t")
        columns_1 = data_file1.loc[0].fillna("")
        testNumber = x[:-4]
        self.label_10.setText(testNumber)
        # Get the Load info for Pre-screening
        if columns_1["Tabbed?"] == "Not Tabbed":
                if columns_1[36] == "Constant Resistor":
                    loadinfo = str(columns_1[17]) + " Ohms for " + str(columns_1[18]) + " Seconds."
                    self.label_4.setText(loadinfo)
                elif columns_1[36] == "Constant Current":
                    loadinfo = str(columns_1[17]) + " mA for " + str(columns_1[18]) + " Seconds."
                    self.label_4.setText(loadinfo)
        else:
            self.label_4.setText("Pre-Tab")
        print("Will this be printed? 6")
        # Get the load infor for Post-Screening
        # Change to Post Table
        if columns_1["Tabbed?"] == "Not Tabbed":
            if not columns_1[27] == "":
                if columns_1[37] == "Constant Resistor":
                    loadinfo = str(columns_1[27]) + " Ohms for " + str(columns_1[28]) + " Seconds."
                    self.label_5.setText(loadinfo)
                elif columns_1[37] == "Constant Current":
                    loadinfo = str(columns_1[27]) + " mA for " + str(columns_1[28]) + " Seconds."
                    self.label_5.setText(loadinfo)
        else:
            if columns_1[36] == "Constant Resistor":
                loadinfo = str(columns_1[17]) + " Ohms for " + str(columns_1[18]) + " Seconds."
                self.label_5.setText(loadinfo)
            elif columns_1[36] == "Constant Current":
                loadinfo = str(columns_1[17]) + " mA for " + str(columns_1[18]) + " Seconds."
                self.label_5.setText(loadinfo)

        # Set Pre-OCV Criteria
        if columns_1["Tabbed?"] == "Not Tabbed":
            pre_OCV_NotTabbed = "%.3f" % round(columns_1[19], 3)
            self.tableWidget.setItem(0, 0, QTableWidgetItem(str(pre_OCV_NotTabbed)))
            # Set Pre-CCV Criteria
            pre_CCV_NotTabbed = "%.3f" % round(columns_1[20], 3)
            self.tableWidget.setItem(0, 1, QTableWidgetItem(str(pre_CCV_NotTabbed)))

        else:
            # Pre-CCV Criteria for tabbing
            pre_CCV_Tabbed = "%.3f" % round(columns_1["Pre-Tab OCV"], 3)
            self.tableWidget.setItem(0, 0, QTableWidgetItem(str(pre_CCV_Tabbed)))

        # Set Post-OCV Criteria check it to see if cell is being tabbed
        if columns_1["Tabbed?"] == "Not Tabbed":
            if columns_1["Profile No."] == 2:
                print("Will this be printed? 5.1")
                post_OCV_Criteria_NotTabbed = "%.3f" % round(columns_1["Profile Two OCV Min"], 3)
                self.tableWidget.setItem(0, 2, QTableWidgetItem(str(post_OCV_Criteria_NotTabbed)))
                print("Will this be printed? 5.2")
                # Set Post-CCV Criteria
                post_CCV_Criteria_NotTabbed = "%.3f" % round(columns_1["Profile Two CCV Min"], 3)
                self.tableWidget.setItem(0, 3, QTableWidgetItem(str(post_CCV_Criteria_NotTabbed)))
                print("Will this be printed? 5.3")
            if columns_1["Section No."] == 2:
                post_OCV_Criteria_NotTabbed = "%.3f" % round(columns_1["Profile One OCV Min"], 3)
                self.tableWidget.setItem(0, 2, QTableWidgetItem(str(post_OCV_Criteria_NotTabbed)))
                # Set Post-CCV Criteria
                post_CCV_Criteria_NotTabbed = "%.3f" % round(columns_1["Profile One CCV Min"], 3)
                self.tableWidget.setItem(0, 3, QTableWidgetItem(str(post_CCV_Criteria_NotTabbed)))

        if columns_1["Tabbed?"] == "Tabbed":
            # Post-OCV Criteria for tabbing
            post_OCV_Criteria = "%.3f" % round(columns_1["Post-Tab OCV"], 3)
            self.tableWidget.setItem(0, 2, QTableWidgetItem(str(post_OCV_Criteria)))
            # Post-CCV Criteria for tabbing
            post_CCV_Criteria = "%.3f" % round(columns_1["Post-Tab CCV"], 3)
            self.tableWidget.setItem(0, 3, QTableWidgetItem(str(post_CCV_Criteria)))
        print("Will this be printed? 5")
        # Set The Max value for Post-OCV
        if not data_file["Post-OCV"].isnull().all():
            post_OCV_Max = "%.3f" % round(data_file.loc[data_file["Post-OCV"].idxmax()]["Post-OCV"], 3)
            self.tableWidget.setItem(1, 2, QTableWidgetItem(str(post_OCV_Max)))
        print("Will this be printed? 6")
        # Set The Max Value for Post-CCV & first check if there is a CCV value
        if not data_file["Post-CCV"].isnull().all():
            post_CCV_Max = "%.3f" % round(data_file.loc[data_file["Post-CCV"].idxmax()]["Post-CCV"], 3)
            self.tableWidget.setItem(1, 3,
                                     QTableWidgetItem(str(post_CCV_Max)))
        print("Will this be printed? 7")
        # Set The Min value for Post-OCV
        if not data_file["Post-OCV"].isnull().all():
            post_OCV_Mix = "%.3f" % round(data_file.loc[data_file["Post-OCV"].idxmin()]["Post-OCV"], 3)
            self.tableWidget.setItem(2, 2,
                                     QTableWidgetItem(str(post_OCV_Mix)))
        print("Will this be printed? 8")
        # Set The Min Value for Post-CCV & first check if there is a CCV value
        if not data_file["Post-CCV"].isnull().all():
            post_CCV_Mix = "%.3f" % round(data_file.loc[data_file["Post-CCV"].idxmin()]["Post-CCV"], 3)
            self.tableWidget.setItem(2, 3,
                                     QTableWidgetItem(str(post_CCV_Mix)))
        # Set the median for Post-OCV
        print("Will this be printed? 3")
        if not data_file["Post-OCV"].isnull().all():
            self.tableWidget.setItem(3, 2, QTableWidgetItem(str("%.3f" % round(data_file["Post-OCV"].median(), 3))))
        # Set the Median for Post-CCV
        if not data_file["Post-CCV"].isnull().all():
            self.tableWidget.setItem(3, 3, QTableWidgetItem(str("%.3f" % round(data_file["Post-CCV"].median(), 3))))
        # set the Mean for Post-OCV
        if not data_file["Post-OCV"].isnull().all():
            self.tableWidget.setItem(4, 2, QTableWidgetItem(str("%.3f" % round(data_file["Post-OCV"].mean(), 3))))
        # Set the Mean for Post-CCV
        if not data_file["Post-CCV"].isnull().all():
            self.tableWidget.setItem(4, 3, QTableWidgetItem(str("%.3f" % round(data_file["Post-CCV"].mean(), 3))))
        # set the S.D for Post-OCV
        if not data_file["Post-OCV"].isnull().all():
            self.tableWidget.setItem(5, 2, QTableWidgetItem(str(round(data_file["Post-OCV"].std(), 5))))
        # set the S.D for Post-CCV
        if not data_file["Post-CCV"].isnull().all():
            self.tableWidget.setItem(5, 3, QTableWidgetItem(str(round(data_file["Post-CCV"].std(), 5))))
        # set the 95% confidence interval Min/Max for Pre-OCV
        print("Will this be printed? 2.5")
        if not data_file["Pre-OCV"].isnull().all():
            data_file1 = data_file["Pre-OCV"].dropna()
            x = self.mean_confidence_interval(data_file1)
            self.tableWidget.setItem(6, 0, QTableWidgetItem(str("%.3f" % round(x[1], 3))))
            self.tableWidget.setItem(7, 0, QTableWidgetItem(str("%.3f" % round(x[2], 3))))
        # set the 95% confidence interval Min/Max for Pre-CCV
        print("Will this be printed? 3.1")
        if not data_file["Pre-CCV"].isnull().all():
            data_file1 = data_file["Pre-CCV"].dropna()
            x = self.mean_confidence_interval(data_file1)
            self.tableWidget.setItem(6, 1, QTableWidgetItem(str("%.3f" % round(x[1], 3))))
            self.tableWidget.setItem(7, 1, QTableWidgetItem(str("%.3f" % round(x[2], 3))))
        # set the 95% confidence interval Min/Max for Post-OCV
        print("Will this be printed? 3.2")
        if not data_file["Post-OCV"].isnull().all():
            data_file1 = data_file["Post-OCV"].dropna()
            x = self.mean_confidence_interval(data_file1)
            self.tableWidget.setItem(6, 2, QTableWidgetItem(str("%.3f" % round(x[1], 3))))
            self.tableWidget.setItem(7, 2, QTableWidgetItem(str("%.3f" % round(x[2], 3))))
        # set the 95% confidence interval Min/Max for Post-CCV
        print("Will this be printed? 3.3")
        if not data_file["Post-CCV"].isnull().all():
            data_file1 = data_file["Post-CCV"].dropna()
            x = self.mean_confidence_interval(data_file1)
            self.tableWidget.setItem(6, 3, QTableWidgetItem(str("%.3f" % round(x[1], 3))))
            self.tableWidget.setItem(7, 3, QTableWidgetItem(str("%.3f" % round(x[2], 3))))
        # set sample outlier pre-OCV
        print("Will this be printed? 2.6")
        if not data_file["Pre-OCV"].isnull().all():
            data_file1 = data_file["Pre-OCV"].dropna()
            sampleOutlier = self.findSampleOutlier(data_file1)
            self.tableWidget.setItem(8, 0, QTableWidgetItem(str(sampleOutlier)))
        # set sample outlier pre-CCV
        print("Will this be printed? 2.7")
        if not data_file["Pre-CCV"].isnull().all():
            data_file1 = data_file["Pre-CCV"].dropna()
            sampleOutlier = self.findSampleOutlier(data_file1)
            self.tableWidget.setItem(8, 1, QTableWidgetItem(str(sampleOutlier)))
        # set sample outlier post-OCV
        print("Will this be printed? 2.8")
        if not data_file["Post-OCV"].isnull().all():
            data_file1 = data_file["Post-OCV"].dropna()
            sampleOutlier = self.findSampleOutlier(data_file1)
            self.tableWidget.setItem(8, 2, QTableWidgetItem(str(sampleOutlier)))
        # set sample outlier post-CCV
        print("Will this be printed? 2.9")
        if not data_file["Post-CCV"].isnull().all():
            data_file1 = data_file["Post-CCV"].dropna()
            sampleOutlier = self.findSampleOutlier(data_file1)
            self.tableWidget.setItem(8, 3, QTableWidgetItem(str(sampleOutlier)))
        # Set the total tested Post-OCV
        print("Will this be printed? 2.1")
        if not data_file["Post-OCV"].isnull().all():
            self.tableWidget.setItem(9, 2, QTableWidgetItem(
                str(round(data_file["Post-OCV"].size - data_file["Post-OCV"].isnull().sum()))))
        # Set the total Tested Post-CCV
        if not data_file["Post-CCV"].isnull().all():
            self.tableWidget.setItem(9, 3, QTableWidgetItem(
                str(round(data_file["Post-CCV"].size - data_file["Post-CCV"].isnull().sum()))))
        # Set The total Failures Post-OCV check to see if the cell is being tabbed?
        print("Will this be printed? 2.2")
        if not data_file["Post-OCV"].isnull().all():
            if columns_1["Tabbed?"] == "Tabbed":
                data_file2 = data_file[data_file["Pre-OCV"] >= columns_1["Pre-Tab OCV"]]
                # pw is how many cell that has PreOCV and PostOCV deference bigger than OCV Tab Tolerance
                allData = round(data_file2["Post-OCV"].dropna(), 3)
                passingData = allData[allData >= columns_1["Post-Tab OCV"]]
                ocvpassingdata = data_file2[data_file2["Post-OCV"] >= columns_1["Post-Tab OCV"]]["Pre-OCV"]
                print("Post-Tab OCV " + str(columns_1["Post-Tab OCV"]))
                pdata = {"OCV": ocvpassingdata, "POCV": passingData}
                df = pd.DataFrame(pdata)
                pw = df[round(abs(df["OCV"] - df["POCV"]), 3) <= columns_1["OCV Tab Tolerance"]]
                failures = len(data_file["Post-OCV"].dropna()) - len(pw)
                self.tableWidget.setItem(10, 2, QTableWidgetItem(str(failures)))
            elif columns_1["Section No."] == 2:
                data_file1 = data_file["Post-OCV"].dropna()
                failures = len(data_file1[data_file1 < columns_1["Profile One OCV Min"]])
                self.tableWidget.setItem(10, 2, QTableWidgetItem(str(failures)))
            else:
                data_file1 = data_file["Post-OCV"].dropna()
                failures = len(data_file1[data_file1 < columns_1["Profile Two OCV Min"]])
                self.tableWidget.setItem(10, 2, QTableWidgetItem(str(failures)))
        # Set The total Failures Pre-CCV check to see if the cell is being tabbed?
        print("Will this be printed? 2.3")
        if not data_file["Post-CCV"].isnull().all():
            if columns_1["Tabbed?"] == "Tabbed":
                data_file1 = data_file["Post-CCV"].dropna()
                failures = len(data_file1[data_file1 < columns_1["Post-Tab CCV"]])


                print(data_file1[data_file1 < columns_1["Post-Tab CCV"]])
                self.tableWidget.setItem(10, 3, QTableWidgetItem(str(failures)))
            elif columns_1["Section No."] == 2:
                data_file1 = data_file["Post-CCV"].dropna()
                failures = len(data_file1[data_file1 < columns_1["Profile One CCV Min"]])
                self.tableWidget.setItem(10, 3, QTableWidgetItem(str(failures)))
            else:
                data_file1 = data_file["Post-CCV"].dropna()
                failures = len(data_file1[data_file1 < columns_1["Profile Two CCV Min"]])
                print(data_file1[data_file1 < columns_1["Profile Two CCV Min"]])
                self.tableWidget.setItem(10, 3, QTableWidgetItem(str(failures)))

        print("Will this be printed? 2")
        '''
        # Set Pre-OCV Criteria
        if columns_1["Tabbed?"] == "Not Tabbed":
            pre_OCV_NotTabbed = "%.3f" % round(columns_1[19], 3)
            self.tableWidget.setItem(0, 0, QTableWidgetItem(str(pre_OCV_NotTabbed)))
            # Set Pre-CCV Criteria
            pre_CCV_NotTabbed = "%.3f" % round(columns_1[20], 3)
            self.tableWidget.setItem(0, 1, QTableWidgetItem(str(pre_CCV_NotTabbed)))

        else:
            # Pre-CCV Criteria for tabbing
            pre_CCV_Tabbed = "%.3f" % round(columns_1["Pre-Tab OCV"], 3)
            self.tableWidget.setItem(0, 0, QTableWidgetItem(str(pre_CCV_Tabbed)))
        '''
        print("Will this be printed? 1")
        # Set The Max value for Pre-OCV
        self.tableWidget.setItem(1, 0, QTableWidgetItem(
            str("%.3f" % round(data_file.loc[data_file["Pre-OCV"].idxmax()]["Pre-OCV"], 3))))
        # Set The Max Value for Pre-CCV & first check if there is a CCV value
        if not data_file["Pre-CCV"].isnull().all():
            self.tableWidget.setItem(1, 1, QTableWidgetItem(
                str("%.3f" % round(data_file.loc[data_file["Pre-CCV"].idxmax()]["Pre-CCV"], 3))))
        print("Will this be printed? 1.1")
        # Set The Min value for Pre-OCV
        self.tableWidget.setItem(2, 0, QTableWidgetItem(
            str("%.3f" % round(data_file.loc[data_file["Pre-OCV"].idxmin()]["Pre-OCV"], 3))))
        # Set The Min Value for Pre-CCV & first check if there is a CCV value
        if not data_file["Pre-CCV"].isnull().all():
            self.tableWidget.setItem(2, 1, QTableWidgetItem(
                str("%.3f" % round(data_file.loc[data_file["Pre-CCV"].idxmin()]["Pre-CCV"], 3))))
        # Set the median for Pre-OCV
        self.tableWidget.setItem(3, 0, QTableWidgetItem(str("%.3f" % round(data_file["Pre-OCV"].median(), 3))))
        # Set the Median for pre-CCV
        print("Will this be printed? 1.2")
        if not data_file["Pre-CCV"].isnull().all():
            self.tableWidget.setItem(3, 1, QTableWidgetItem(str("%.3f" % round(data_file["Pre-CCV"].median(), 3))))
        # set the Mean for Pre-OCV
        self.tableWidget.setItem(4, 0, QTableWidgetItem(str("%.3f" % round(data_file["Pre-OCV"].mean(), 3))))
        # Set the Mean for pre-CCV
        print("Will this be printed? 1.3")
        if not data_file["Pre-CCV"].isnull().all():
            self.tableWidget.setItem(4, 1, QTableWidgetItem(str("%.3f" % round(data_file["Pre-CCV"].mean(), 3))))
        # set the S.D for Pre-OCV
        self.tableWidget.setItem(5, 0, QTableWidgetItem(str(round(data_file["Pre-OCV"].std(), 5))))
        # set the S.D for Pre-CCV
        if not data_file["Pre-CCV"].isnull().all():
            self.tableWidget.setItem(5, 1, QTableWidgetItem(str(round(data_file["Pre-CCV"].std(), 5))))
        print("Will this be printed? 1.4")
        # set the Range Min Pre-OCV

        # set the Range Min Pre-CCV

        # set the Range Max Pre-OCV

        # set the Range Max Pre-CCV

        # Set the total tested Pre-OCV
        self.tableWidget.setItem(9, 0, QTableWidgetItem(
            str(round(data_file["Pre-OCV"].size - data_file["Pre-OCV"].isnull().sum()))))
        # Set the total Tested Pre-CCV
        if not data_file["Pre-CCV"].isnull().all():
            self.tableWidget.setItem(9, 1, QTableWidgetItem(
                str(round(data_file["Pre-CCV"].size - data_file["Pre-CCV"].isnull().sum()))))
        # Set The total Failures Pre-OCV check to see if the cell is being tabbed?
        print("Will this be printed? 9.9")
        if not data_file["Pre-OCV"].isnull().all():
            if columns_1["Tabbed?"] == "Tabbed":
                data_file2 = data_file[data_file["Pre-OCV"] >= columns_1["Pre-Tab OCV"]]
                # pw is how many cell that has PreOCV and PostOCV deference bigger than OCV Tab Tolerance
                allData = round(data_file2["Post-OCV"].dropna(), 3)
                passingData = allData[allData >= columns_1["Post-Tab OCV"]]
                ocvpassingdata = data_file2[data_file2["Post-OCV"] >= columns_1["Post-Tab OCV"]]["Pre-OCV"]
                print("Post-Tab OCV " + str(columns_1["Post-Tab OCV"]))
                pdata = {"OCV": ocvpassingdata, "POCV": passingData}
                df = pd.DataFrame(pdata)
                pw = df[round(abs(df["OCV"] - df["POCV"]), 3) <= columns_1["OCV Tab Tolerance"]]
                failures = len(data_file["Pre-OCV"].dropna()) - len(pw)
                self.tableWidget.setItem(10, 0, QTableWidgetItem(str(failures)))
            else:
                data_file1 = data_file["Pre-OCV"].dropna()
                failures = len(data_file1[data_file1 < columns_1["Profile One OCV Min"]])
                self.tableWidget.setItem(10, 0, QTableWidgetItem(str(failures)))
        # Set The total Failures Pre-CCV check to see if the cell is being tabbed?
        if not data_file["Pre-CCV"].isnull().all():
            if columns_1["Tabbed?"] == "Not Tabbed":
                data_file1 = data_file["Pre-CCV"].dropna()
                failures = len(data_file1[data_file1 < columns_1["Profile One CCV Min"]])
                self.tableWidget.setItem(10, 1, QTableWidgetItem(str(failures)))

        # setting labels
        if columns_1["Tabbed?"] == "Tabbed":
            self.tableWidget.setHorizontalHeaderLabels(["Pre-Tab OCV", "", "Post-Tab OCV", "Post-Tab CCV"])
            self.label_7.setText("Pre-Tab OCV")
            self.label_11.setText("Post-Tab OCV")
            self.label_12.setText("Post-Tab CCV")
        elif columns_1["Profile No."] == 2:
            self.tableWidget.setHorizontalHeaderLabels(["Profile One OCV", "Profile One CCV", "Profile Two OCV", "Profile Two CCV"])
            self.label_7.setText("Profile One OCV")
            self.label_8.setText("Profile One CCV")
            self.label_11.setText("Profile Two OCV")
            self.label_12.setText("Profile Two CCV")
        elif columns_1["Section No."] == 2:
            self.tableWidget.setHorizontalHeaderLabels(
                ["Section One OCV", "Section One CCV", "Section Two OCV", "Section Two CCV"])
            self.label_7.setText("Section One OCV")
            self.label_8.setText("Section One CCV")
            self.label_11.setText("Section Two OCV")
            self.label_12.setText("Section Two CCV")
            if columns_1[36] == "Constant Resistor":
                loadinfo = str(columns_1["Profile One Values"]) + " Ohms for " + str(columns_1["Profile One Timer"]) + " Seconds."
                self.label_5.setText(loadinfo)
            elif columns_1[36] == "Constant Current":
                loadinfo = str(columns_1["Profile One Values"]) + " mA for " + str(columns_1["Profile One Timer"]) + " Seconds."
                self.label_5.setText(loadinfo)
        else:
            self.tableWidget.setHorizontalHeaderLabels(
                ["OCV", "CCV", "", ""])
            self.label_7.setText("OCV")
            self.label_8.setText("CCV")

        print("Will this be printed? 1")
        # Show the Pre table_OCV first find the BinWith of the histgram
        if not data_file["Pre-OCV"].isnull().all():
            data_file["Pre-OCV"] = round(data_file["Pre-OCV"], 3)
            print(data_file.loc[data_file["Pre-OCV"].idxmax()]["Pre-OCV"])
            print(data_file.loc[data_file["Pre-OCV"].idxmin()]["Pre-OCV"])
            MaxtoMin = round(data_file.loc[data_file["Pre-OCV"].idxmax()]["Pre-OCV"] -
                             data_file.loc[data_file["Pre-OCV"].idxmin()]["Pre-OCV"], 3)
            HowManyBin = int(round(math.sqrt(data_file["Pre-OCV"].size - data_file["Pre-OCV"].isnull().sum()))) + 2
            histgram_BinWith = round((MaxtoMin / HowManyBin), 3) + 0.001
            # if the sample size is too small, this will make sure the binWith is aleast 0.001
            if histgram_BinWith < 0.001:
                histgram_BinWith = 0.001
            lowest_point = data_file.loc[data_file["Pre-OCV"].idxmin()]["Pre-OCV"] - histgram_BinWith
            eachBin = data_file.loc[data_file["Pre-OCV"].idxmax()]["Pre-OCV"]
            self.tableWidget_4.setRowCount(HowManyBin)
            pre_highIndex = 0
            while eachBin >= lowest_point:
                lowBin = round(eachBin - round(histgram_BinWith, 3), 3)
                highBin = round(eachBin, 3)
                if round(eachBin - round(histgram_BinWith, 3), 3) <= lowest_point:
                    self.tableWidget_4.setItem(pre_highIndex, 0, QTableWidgetItem(str("%.3f" % round(highBin, 3))))
                    self.tableWidget_4.setItem(pre_highIndex, 1, QTableWidgetItem(str("%.3f" % round(lowBin, 3))))
                    self.tableWidget_4.setItem(pre_highIndex, 2, QTableWidgetItem(
                        str(data_file["Pre-OCV"][data_file["Pre-OCV"].between(lowBin, highBin)].size)))
                else:
                    self.tableWidget_4.setItem(pre_highIndex, 0, QTableWidgetItem(str("%.3f" % round(highBin, 3))))
                    self.tableWidget_4.setItem(pre_highIndex, 1, QTableWidgetItem(str("%.3f" % round(lowBin, 3))))
                    self.tableWidget_4.setItem(pre_highIndex, 2, QTableWidgetItem(
                        str(data_file["Pre-OCV"][data_file["Pre-OCV"].between(lowBin + 0.001, highBin)].size)))
                eachBin = round(eachBin - round(histgram_BinWith, 3), 3)
                pre_highIndex = pre_highIndex + 1
        print("Will this be printed? 11.1")
        # Show the Pre table_CCV first find the BinWith of the histgram
        if not data_file["Pre-CCV"].isnull().all():
            data_file["Pre-CCV"] = round(data_file["Pre-CCV"], 3)
            MaxtoMin = data_file.loc[data_file["Pre-CCV"].idxmax()]["Pre-CCV"] - \
                       data_file.loc[data_file["Pre-CCV"].idxmin()]["Pre-CCV"]
            HowManyBin = int(round(math.sqrt(data_file["Pre-CCV"].size - data_file["Pre-CCV"].isnull().sum()))) + 2
            histgram_BinWith = round((MaxtoMin / HowManyBin), 3) + 0.001
            #if the sample size is too small, this will make sure the binWith is aleast 0.001
            if histgram_BinWith < 0.001:
                histgram_BinWith = 0.001
            lowest_point = data_file.loc[data_file["Pre-CCV"].idxmin()]["Pre-CCV"]
            eachBin = data_file.loc[data_file["Pre-CCV"].idxmax()]["Pre-CCV"]
            self.tableWidget_2.setRowCount(HowManyBin)
            pre_highIndex = 0
            while eachBin > lowest_point:
                lowBin = round(eachBin - round(histgram_BinWith, 3), 3)
                highBin = round(eachBin, 3)
                if round(eachBin - round(histgram_BinWith, 3), 3) <= lowest_point:
                    self.tableWidget_2.setItem(pre_highIndex, 0, QTableWidgetItem(str("%.3f" % round(highBin, 3))))
                    self.tableWidget_2.setItem(pre_highIndex, 1, QTableWidgetItem(str("%.3f" % round(lowBin, 3))))
                    self.tableWidget_2.setItem(pre_highIndex, 2, QTableWidgetItem(
                        str(data_file["Pre-CCV"][data_file["Pre-CCV"].between(lowBin, highBin)].size)))
                else:
                    self.tableWidget_2.setItem(pre_highIndex, 0, QTableWidgetItem(str("%.3f" % round(highBin, 3))))
                    self.tableWidget_2.setItem(pre_highIndex, 1, QTableWidgetItem(str("%.3f" % round(lowBin, 3))))
                    self.tableWidget_2.setItem(pre_highIndex, 2, QTableWidgetItem(
                        str(data_file["Pre-CCV"][data_file["Pre-CCV"].between(lowBin+0.001, highBin)].size)))
                eachBin = round(eachBin-round(histgram_BinWith, 3), 3)
                pre_highIndex = pre_highIndex+1
        #self.tableWidget_2.sortItems(2, order=Qt.DescendingOrder)
        print("Will this be printed? 11.2")
        # Show the Post table_OCV first find the BinWith of the histgram
        if not data_file["Post-OCV"].isnull().all():

            data_file["Post-OCV"] = round(data_file["Post-OCV"], 3)
            MaxtoMin = data_file.loc[data_file["Post-OCV"].idxmax()]["Post-OCV"] - \
                       data_file.loc[data_file["Post-OCV"].idxmin()]["Post-OCV"]
            HowManyBin = int(
                round(math.sqrt(data_file["Post-OCV"].size - data_file["Post-OCV"].isnull().sum()))) + 2
            histgram_BinWith = round((MaxtoMin / HowManyBin), 3) + 0.001
            # if the sample size is too small, this will make sure the binWith is aleast 0.001
            if histgram_BinWith < 0.001:
                histgram_BinWith = 0.001
            lowest_point = data_file.loc[data_file["Post-OCV"].idxmin()]["Post-OCV"]
            eachBin = data_file.loc[data_file["Post-OCV"].idxmax()]["Post-OCV"]
            self.tableWidget_5.setRowCount(HowManyBin)
            pre_highIndex = 0
            while eachBin > lowest_point:
                lowBin = round(eachBin - round(histgram_BinWith, 3), 3)
                highBin = round(eachBin, 3)
                if round(eachBin - round(histgram_BinWith, 3), 3) <= lowest_point:
                    self.tableWidget_5.setItem(pre_highIndex, 0, QTableWidgetItem(str("%.3f" % round(highBin, 3))))
                    self.tableWidget_5.setItem(pre_highIndex, 1, QTableWidgetItem(str("%.3f" % round(lowBin, 3))))
                    self.tableWidget_5.setItem(pre_highIndex, 2, QTableWidgetItem(
                        str(data_file["Post-OCV"][data_file["Post-OCV"].between(lowBin, highBin)].size)))
                else:
                    self.tableWidget_5.setItem(pre_highIndex, 0, QTableWidgetItem(str("%.3f" % round(highBin, 3))))
                    self.tableWidget_5.setItem(pre_highIndex, 1, QTableWidgetItem(str("%.3f" % round(lowBin, 3))))
                    self.tableWidget_5.setItem(pre_highIndex, 2, QTableWidgetItem(
                        str(data_file["Post-OCV"][
                                data_file["Post-OCV"].between(lowBin + 0.001, highBin)].size)))
                eachBin = round(eachBin - round(histgram_BinWith, 3), 3)
                pre_highIndex = pre_highIndex + 1

        print("Will this be printed? 11.3")
        # Show the Post table first find the BinWith of the histgram
        if not data_file["Post-CCV"].isnull().all():

            data_file["Post-CCV"] = round(data_file["Post-CCV"], 3)
            MaxtoMin = data_file.loc[data_file["Post-CCV"].idxmax()]["Post-CCV"] - \
                       data_file.loc[data_file["Post-CCV"].idxmin()]["Post-CCV"]
            HowManyBin = int(round(math.sqrt(data_file["Post-CCV"].size - data_file["Post-CCV"].isnull().sum()))) + 2
            histgram_BinWith = round((MaxtoMin / HowManyBin), 3) + 0.001
            # if the sample size is too small, this will make sure the binWith is aleast 0.001
            if histgram_BinWith < 0.001:
                histgram_BinWith = 0.001
            lowest_point = data_file.loc[data_file["Post-CCV"].idxmin()]["Post-CCV"]
            eachBin = data_file.loc[data_file["Post-CCV"].idxmax()]["Post-CCV"]
            self.tableWidget_3.setRowCount(HowManyBin)
            pre_highIndex = 0
            while eachBin > lowest_point:
                lowBin = round(eachBin - round(histgram_BinWith, 3), 3)
                highBin = round(eachBin, 3)
                if round(eachBin - round(histgram_BinWith, 3), 3) <= lowest_point:
                    self.tableWidget_3.setItem(pre_highIndex, 0, QTableWidgetItem(str("%.3f" % round(highBin, 3))))
                    self.tableWidget_3.setItem(pre_highIndex, 1, QTableWidgetItem(str("%.3f" % round(lowBin, 3))))
                    self.tableWidget_3.setItem(pre_highIndex, 2, QTableWidgetItem(
                        str(data_file["Post-CCV"][data_file["Post-CCV"].between(lowBin, highBin)].size)))
                else:
                    self.tableWidget_3.setItem(pre_highIndex, 0, QTableWidgetItem(str("%.3f" % round(highBin, 3))))
                    self.tableWidget_3.setItem(pre_highIndex, 1, QTableWidgetItem(str("%.3f" % round(lowBin, 3))))
                    self.tableWidget_3.setItem(pre_highIndex, 2, QTableWidgetItem(
                        str(data_file["Post-CCV"][data_file["Post-CCV"].between(lowBin+0.001, highBin)].size)))
                eachBin = round(eachBin-round(histgram_BinWith, 3), 3)
                pre_highIndex = pre_highIndex + 1
        self.screening_Report_wordGenerator(testNumber, columns_1, self.tableWidget, self.tableWidget_4,
                                            self.tableWidget_2, self.tableWidget_5, self.tableWidget_3, self.label_4, self.label_5)


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    qt_app = Report_StatisticsTable_WithFunctions()
    qt_app.getTestNumber2("14575A00.txt")
    qt_app.show()
    app.exec_()






