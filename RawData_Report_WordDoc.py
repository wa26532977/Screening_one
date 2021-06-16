import pandas as pd
import numpy as np
import docx
import os
import datetime
from docx import Document
from docx.shared import Pt
import sys

pd.options.display.max_columns = 999
pd.options.display.max_rows = 999
pd.set_option("display.precision", 6)


def RawData_report_wordDoc(x):
    # import the template
    doc = docx.Document(os.path.dirname(sys.argv[0]) + r"\\Report_Word\\Doc Template\\RawData-template.docx")

    # i = 0
    # show the doc
    # for text in doc.paragraphs:
    #     run_i = 0
    #     for runs in text.runs:
    #         print(runs.text)
    #         print("This is " + str(i) + " lin and " + str(run_i) + " runs")
    #         run_i = run_i + 1
    #     i = i + 1

    # get the template info
    path_template = os.path.dirname(sys.argv[0]) + r"\\Screening_Template\\" + x
    data_file1 = pd.read_csv(path_template, sep="\t")
    columns_1 = data_file1.loc[0].fillna("")
    # get the testing data
    path_datafile = os.path.dirname(sys.argv[0]) + r"\\Screening_Data\\" + x
    data_file = pd.read_csv(path_datafile, sep="\t")
    data_file = data_file.fillna("")
    # get the outlier data
    path_outlier_data = os.path.dirname(sys.argv[0]) + r"\\Report_word_Outlier\\" + x[0: -4] + r"outlier.txt"
    outlier_data = pd.read_csv(path_outlier_data, sep="\t")
    outlier_data = outlier_data.replace(np.nan, '', regex=True)
    print(outlier_data)

    page_number = len(data_file["Barcode"]) // 24
    if len(data_file["Barcode"]) % 24 != 0:
        page_number = page_number + 1

    indexCount = 0
    current_page = 1
    testing = 0

    while indexCount < len(data_file["Barcode"]):
        # if indexCount == 0:
        #    doc.paragraphs[0].runs[5].text = str(1)
        # elif indexCount % 24 == 0:
        current_page += 1
        #   doc.paragraphs[0].runs[5].text = str(current_page)
        # add data into the table
        table = doc.tables[0]
        if indexCount == 0:
            # set the Test Number
            doc.paragraphs[1].runs[1].text = str(x[:-4])
            # set the chemistry
            doc.paragraphs[1].runs[9].text = str(columns_1["Chemistry"]).translate(str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉"))
            # Set the sample name
            doc.paragraphs[2].runs[2].text = str(columns_1["Cell Name"])
            # set load
            if columns_1["Profile No."] == 2:
                # set the table index
                table.rows[0].cells[3].text = "Profile 1 OCV"
                table.rows[0].cells[4].text = "Profile 1 CCV"
                table.rows[0].cells[5].text = "Profile 2 OCV"
                table.rows[0].cells[6].text = "Profile 2 CCV"

                doc.paragraphs[2].runs[9].text = str(columns_1["Profile One Values"]) + "/" + str(
                    columns_1["Profile Two Value"])
                # timer
                doc.paragraphs[3].runs[9].text = str(columns_1["Profile One Timer"]) + "/" + str(
                    columns_1["Profile Two Timer"])
            elif columns_1["Section No."] == 2:
                # set the table index
                table.rows[0].cells[3].text = "Section 1 OCV"
                table.rows[0].cells[4].text = "Section 1 CCV"
                table.rows[0].cells[5].text = "Section 2 OCV"
                table.rows[0].cells[6].text = "Section 2 CCV"

                doc.paragraphs[2].runs[9].text = str(columns_1["Profile One Values"]) + "/" + str(
                    columns_1["Profile One Values"])
                doc.paragraphs[3].runs[9].text = str(columns_1["Profile One Timer"]) + "/" + str(
                    columns_1["Profile One Timer"])
            elif columns_1["Tabbed?"] == "Tabbed":
                # set the table index
                table.rows[0].cells[3].text = "Pre-Tab OCV"
                table.rows[0].cells[4].text = ""
                table.rows[0].cells[5].text = "Post-Tab OCV"
                table.rows[0].cells[6].text = "Post-Tab CCV"

                doc.paragraphs[2].runs[9].text = " /" + str(columns_1["Profile One Values"])
                doc.paragraphs[3].runs[9].text = " /" + str(columns_1["Profile One Timer"])
            else:
                # set the table index
                table.rows[0].cells[5].text = ""
                table.rows[0].cells[6].text = ""

                doc.paragraphs[2].runs[9].text = str(columns_1["Profile One Values"])
                doc.paragraphs[3].runs[9].text = str(columns_1["Profile One Timer"])

            # the set the unit for the load
            if columns_1["Profile One Type"] == "Constant Resistor":
                doc.paragraphs[2].runs[11].text = "Ω"
            else:
                doc.paragraphs[2].runs[11].text = "mV"

        # add new row to the table
        cell_info = table.add_row()
        # set the data to the table
        cell_counter = 0

        for cell in cell_info.cells:
            if cell_counter <= 1:
                if data_file.iloc[indexCount][cell_counter] != "":
                    cell.text = str(int(round(data_file.iloc[indexCount][cell_counter])))
                else:
                    cell.text = str(data_file.iloc[indexCount][cell_counter])
            elif cell_counter == 2:
                cell.text = str(columns_1["Mfg Date"])
            # the visual Inspection
            elif cell_counter == 7:
                cell.text = outlier_data["Outlier"][indexCount]
            else:
                if isinstance(outlier_data.iloc[indexCount][cell_counter-2], float):
                    cell.text = "%.3f" % (round(outlier_data.iloc[indexCount][cell_counter-2], 3))
                else:
                    cell.text = str(outlier_data.iloc[indexCount][cell_counter-2])
            cell_counter += 1

        indexCount = indexCount + 1

    doc.add_paragraph("Key: ^ for Tab Tolerance Fail(T), ! for Criteria Fall(F), * for outlier(OH for High)(OL for Low)")

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(9)

    doc.save(os.path.dirname(sys.argv[0]) + r"\\Report_Word\\" + x[0:-4] + "RawData.docx")


if __name__ == '__main__':
    RawData_report_wordDoc("peterTest12345.txt")
