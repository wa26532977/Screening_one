# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'Testing_dialog_peter.ui'
#
# Created by: PyQt5 UI code generator 5.11.3
#
# WARNING! All changes made in this file will be lost!
#
# from PyQt5 import QtCore, QtGui, QtWidgets
#
# class Ui_Dialog(object):
#     def setupUi(self, Dialog):
#         Dialog.setObjectName("Dialog")
#         Dialog.resize(400, 300)
#         self.buttonBox = QtWidgets.QDialogButtonBox(Dialog)
#         self.buttonBox.setGeometry(QtCore.QRect(30, 240, 341, 32))
#         self.buttonBox.setOrientation(QtCore.Qt.Horizontal)
#         self.buttonBox.setStandardButtons(QtWidgets.QDialogButtonBox.Cancel|QtWidgets.QDialogButtonBox.Ok)
#         self.buttonBox.setObjectName("buttonBox")
#         self.label = QtWidgets.QLabel(Dialog)
#         self.label.setGeometry(QtCore.QRect(50, 40, 241, 91))
#         font = QtGui.QFont()
#         font.setPointSize(20)
#         self.label.setFont(font)
#         self.label.setObjectName("label")
#
#         self.retranslateUi(Dialog)
#         self.buttonBox.accepted.connect(Dialog.accept)
#         self.buttonBox.rejected.connect(Dialog.reject)
#         QtCore.QMetaObject.connectSlotsByName(Dialog)
#
#     def retranslateUi(self, Dialog):
#         _translate = QtCore.QCoreApplication.translate
#         Dialog.setWindowTitle(_translate("Dialog", "Dialog"))
#         self.label.setText(_translate("Dialog", "Happy Peter"))
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Inches
from docxcompose.composer import Composer
import os
# # import pandas as pd
# path_data = r"C:\Users\wangp.BTC\PycharmProjects\BTC-Work\Screening_System_PyQt5\Report_Word\\"
# # outlier_data = pd.read_csv(path_data, sep="\t")
# # print(outlier_data.iloc[0][1])
# import docx
# from docxtpl import DocxTemplate
#
# doc = docx.Document(os.getcwd() + r"\\Report_Word\\Doc Template\\Static-template-PreTab-New.docx")
# doc.paragraphs[13].runs[3].text = str("table_one")
# print(doc.paragraphs[14].runs[5].text, doc.paragraphs[14].runs[6].text)
# f1 = path_data + r"14575A00.docx"
# f_note = path_data + r"Note-14575A00.docx"
# f2 = path_data + r"14575A00Pre-StaticReport.docx"
# f3 = path_data + r"14575A00Post-StaticReport.docx"
# f4 = path_data + r"14575A00RawData.docx"


# doc.save(saving_dict)
#
# doc_1 = DocxTemplate(saving_dict)
# context = {'c_1': "test", "c_4": "other"}
# context['c_2'] = "goodcode"
# doc_1.render(context)
# # doc_1.save("Test_generated_doc.docx")


# master = Document(f1)
# composer = Composer(master)
# note = Document(f_note)
# doc1 = Document(f2)
# doc2 = Document(f3)
# doc3 = Document(f4)
# composer.append(note)
# composer.append(doc1)
# composer.append(doc2)
# composer.append(doc3)
# composer.save("Testing_123.docx")
# from os import listdir
# from os.path import isfile, join
# import glob
import sys

# import os
# import pandas as pd
#
# pd.options.display.max_columns = 999
# pd.options.display.max_rows = 999
# pd.set_option("display.precision", 6)

# print(os.path.dirname(sys.argv[0]))
# print(1)
# print(os.getcwd())

# search_dir = r"C:\Users\wangp.BTC\PycharmProjects\BTC-Work\Screening_System_PyQt5\Screening_Data\\"
# # remove anything from the list that is not a file (directories, symlinks)
# # of files (presumably not including directories)
# files = list(filter(os.path.isfile, glob.glob(search_dir + "*")))
# files.sort(key=lambda x: os.path.getmtime(x), reverse=True)
# print(files)
# for item in files:
#     path_datafile = item
#     data_file = pd.read_csv(path_datafile, sep="\t")
#     # print(data_file["Barcode"] == 455152)
#     if 455152 in data_file["Barcode"].values:
#         print(item)
#         print(data_file[data_file["Barcode"] == 455152])
#         break

# try_data = [('1.7', '0.0 A', 0.13147), ('1.679', '0.0997 A', 0.37899), ('1.678', '0.0997 A', 0.51106),
#             ('1.677', '0.0997 A', 0.74405), ('1.676', '0.0997 A', 0.88745), ('1.676', '0.0998 A', 1.11443),
#             ('1.675', '0.0998 A', 1.29228), ('1.675', '0.0997 A', 1.51983), ('1.674', '0.0997 A', 1.6838),
#             ('1.673', '0.0997 A', 1.90978), ('1.673', '0.0997 A', 2.05379), ('1.672', '0.0997 A', 2.27743),
#             ('1.672', '0.0997 A', 2.41719), ('1.671', '0.0998 A', 2.5555), ('1.67', '0.0997 A', 2.80457),
#             ('1.67', '0.0997 A', 2.96784), ('1.67', '0.0997 A', 3.17905), ('1.669', '0.0997 A', 3.33342),
#             ('1.669', '0.0997 A', 3.54949), ('1.668', '0.0997 A', 3.71654), ('1.667', '0.0997 A', 3.96579),
#             ('1.667', '0.0997 A', 4.12527), ('1.666', '0.0997 A', 4.33846), ('1.666', '0.0998 A', 4.49558),
#             ('1.666', '0.0998 A', 4.71795), ('1.665', '0.0997 A', 4.8773), ('1.665', '0.0997 A', 5.12753),
#             ('1.665', '0.0861 A', 5.40291), ('1.684', '0.0 A', 5.56935), ('1.685', '0.0 A', 5.70194),
#             ('1.686', '0.0 A', 5.91985), ('1.686', '0.0 A', 6.08108)]
#
# import numpy as np
# dir_path = os.path.dirname(sys.argv[0])
# path_fast_data = dir_path + r"\\Report_word_Outlier\\" + "14872B00outlier.txt"
# fast_data_file = pd.read_csv(path_fast_data, sep="\t")
# outlier_data = fast_data_file.replace(np.nan, '', regex=True)
# # # print(fast_data_file[fast_data_file["Barcode"] == 1111]["Pre"].values[0])
# # # print(fast_data_file[fast_data_file["Barcode"] == 1111]["Pre"].values[0])
# # # print(fast_data_file[fast_data_file["Barcode"] == 2222]["Pre"].values[0])
# # print(fast_data_file[fast_data_file["Barcode"] == 1111])
# # print(fast_data_file[fast_data_file["Barcode"] == 1111]["Pre"])
# # value = fast_data_file[fast_data_file["Barcode"] == 1111]["Post"].values[0]
# # print(pd.isna(value))
# print(outlier_data)
# print(outlier_data.iloc[0][3])
# print(outlier_data.loc[0][3])
# print(type(outlier_data.loc[0][2]))
# print(type(outlier_data.iloc[0][2]))
# for i in range(3):
#     for c in range(5):
#         if isinstance(outlier_data.iloc[i][c], float):
#             print(f"float!!! = {outlier_data.iloc[i][c]}")
#         else:
#             print(outlier_data.iloc[i][c])

# report_doc = Document(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\14872B00.docx")
# document = Document(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\14872B00.docx")

# document = Document()
# section = document.sections[-1]  # Sections object
# new_width, new_height = section.page_height, section.page_width
# section.orientation = WD_ORIENT.LANDSCAPE
# section.page_width = new_width
# section.page_height = new_height
# doc = document.add_paragraph().add_run()
# doc.add_picture("Pre OCV 95% confidence interval.png", width=Inches(9))
#
# print(section.start_type, section.orientation)
# document.save(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\" + "thisISTest" + '.docx')

selected_file = "14872C00"+".txt"
testName = selected_file[0:-4]
files = [testName+'.docx', "Note-" + testName + ".docx", testName + "Pre-StaticReport.docx"]
if os.path.exists(os.path.dirname(sys.argv[0]) + r"\\Report_Word\\" + testName + "Post-StaticReport.docx"):
    files.append(testName + "Post-StaticReport.docx")
files.append(testName + "RawData.docx")
files.append(testName + "_Graph.docx")

# files.pop(2)
print(files)

merged_document = Document()
'''
for index, file in enumerate(files):
    sub_doc = Document(os.path.dirname(sys.argv[0]) + r"\\Report_Word\\" + file)
    for element in sub_doc.element.body:
        print(element)
        merged_document.element.body.append(element)
merged_document.save(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\1_" + testName + '.docx')

merged_document = Document(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\1_" + testName + '.docx')
# add pictures
for section in merged_document.sections:
    print(section.start_type, section.orientation)
section = merged_document.sections[-1]  # Sections object
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height
print("Second sections:")
for section in merged_document.sections:
    print(section.start_type, section.orientation)
merged_document.add_paragraph("this works?")
'''

# merged_document.add_picture("OCV.png", width=Inches(9))
# doc = merged_document.add_paragraph().add_run()
# doc.add_picture("Pre OCV 95% confidence interval.png", width=Inches(9))

# doc.add_picture("OCV.png", width=Inches(9))
# doc.add_picture("Pre OCV WITHIN SD RANGE.png", width=Inches(9))
# doc.add_picture("Pre OCV 95% confidence interval.png", width=Inches(9))
#
# doc.add_picture("OCV 2.png", width=Inches(9))
# doc.add_picture("Post OCV WITHIN SD RANGE.png", width=Inches(9))
# doc.add_picture("CCV 2.png", width=Inches(9))
# doc.add_picture("Post CCV WITHIN SD RANGE.png", width=Inches(9))
# doc.add_picture("Post OCV 95% confidence interval.png", width=Inches(9))
# doc.add_picture("Post CCV 95% confidence interval.png", width=Inches(9))
# doc.add_picture("OCV-CCV 2.png", width=Inches(9))
# doc.add_picture("Post OCV-CCV WITHIN SD RANGE.png", width=Inches(9))


path_name = os.path.dirname(sys.argv[0]) + r"\\Report_Word\\"
master = Document(path_name + testName + '.docx')
composer = Composer(master)
for item in files:
    doc_add = Document(path_name + item)
    composer.append(doc_add)

composer.save(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\" + testName + '.docx')
# merged_document.save(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\" + testName + '.docx')
os.startfile(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\" + testName + '.docx')


