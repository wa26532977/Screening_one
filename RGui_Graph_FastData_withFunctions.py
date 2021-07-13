import os
import docx
from docx.shared import Inches
import pandas as pd
import sys
from PyQt5.QtWidgets import QDialog, QFileDialog, QMessageBox
from PyQt5 import QtWidgets
from PyQt5.uic import loadUi
import ast
import matplotlib.pyplot as plt
from PyQt5.QtGui import *
import numpy as np


pd.options.display.max_columns = 999
pd.options.display.max_rows = 999
pd.set_option("display.precision", 6)


class GraphFastDataWithFunction(QDialog):
    def __init__(self, test_name):
        super(GraphFastDataWithFunction, self).__init__()
        loadUi("RGui_Graph_FastData.ui", self)
        self.test_name = test_name
        self.label_3.setText(self.test_name)
        self.dir_path = os.path.dirname(sys.argv[0])
        self.path_fast_data = self.dir_path + r"\\Fast_data_collection\\" + self.test_name
        self.fast_data_file = pd.read_csv(self.path_fast_data, sep="\t")
        self.path_start_data = self.dir_path + r"\\Screening_Data\\" + self.test_name
        self.start_data_file = pd.read_csv(self.path_start_data, sep="\t")
        self.data_set = []
        if self.fast_data_file["Pre"].dropna().size == 0:
            self.radioButton.setCheckable(False)
            self.radioButton_2.setChecked(True)
        self.populate_list_widget()
        self.x_axis_max = 0
        self.y_axis_min = 0
        self.y_axis_max = 0
        self.pushButton_2.clicked.connect(self.graph_one)
        self.pushButton_3.clicked.connect(self.graph_all)
        self.pushButton_4.clicked.connect(self.set_y_axis)
        self.pushButton_5.clicked.connect(self.reset_y_axis)
        self.pushButton.clicked.connect(self.print_graph)

    def print_graph(self):
        doc = docx.Document(self.dir_path + r'\\Report_Word\\Doc Template\\FastData_Graph.docx')
        doc.add_picture("fast_data_test.png", width=Inches(10))
        print(self.data_set)
        if self.data_set:
            doc.add_paragraph(str(self.data_set))
        doc.save(self.dir_path + r"\\Final_Report\\fast_data_" + self.test_name + '.docx')
        os.startfile(self.dir_path + r"\\Final_Report\\fast_data_" + self.test_name + '.docx')

    def set_test_name(self, test_name):
        self.test_name = test_name
        self.path_fast_data = self.dir_path + r"\\Fast_data_collection\\" + self.test_name
        self.fast_data_file = pd.read_csv(self.path_fast_data, sep="\t")
        self.path_start_data = self.dir_path + r"\\Screening_Data\\" + self.test_name
        self.start_data_file = pd.read_csv(self.path_start_data, sep="\t")

    def reset_y_axis(self):
        self.y_axis_min = 0
        self.y_axis_max = 0
        self.lineEdit.setText("")
        self.lineEdit_2.setText("")
        self.graph_one()

    def set_y_axis(self):
        if self.lineEdit.text() != '':
            self.y_axis_min = float(self.lineEdit.text())
        if self.lineEdit_2.text() != "":
            self.y_axis_max = float(self.lineEdit_2.text())
        self.graph_one()

    def graph_all(self):
        profile_select = "Post"
        data_dict = {}
        if self.radioButton.isChecked():
            profile_select = "Pre"

        select_data = self.fast_data_file[profile_select][0]
        # if the select data is nan give warning
        if pd.isna(select_data):
            msgbox = QtWidgets.QMessageBox(self)
            msgbox.setText(f"There is no {profile_select} test data!")
            msgbox.exec()
            return

        for index, item in self.fast_data_file.iterrows():
            if not pd.isna(item[profile_select]):
                select_data_1 = ast.literal_eval(item[profile_select])
                voltage_list = []
                time_list = []
                for data in select_data_1:
                    time_check = select_data_1[1][2]
                    if not voltage_list:
                        init_data = self.start_data_file[self.start_data_file["Barcode"] == int(item["Barcode"])][profile_select+'-OCV'].values[0]
                        voltage_list.append(init_data)
                        time_list.append(-1)
                        voltage_list.append(init_data)
                        time_list.append(-0.1)
                    else:
                        voltage_list.append(float(data[0]))
                        time_list.append(data[2]-time_check)
                    data_dict[item['Barcode']] = [voltage_list, time_list]

        # graph all the plot
        fig, ax = plt.subplots()
        plt.title("Fast Data plot " + self.test_name[:-4] + " All Cells at " + profile_select , fontsize=22)
        fig.set_size_inches(23, 18)
        plt.rcParams.update({"font.size": 22})
        plt.tick_params(labelsize=22)
        for key, value in data_dict.items():
            plt.plot(value[1], value[0], label=key)
        plt.legend()
        plt.xlabel('time (S)', fontsize=22)
        plt.ylabel('Voltage (V)', fontsize=22)
        # y-axis lim
        if self.y_axis_min != 0:
            plt.ylim([self.y_axis_min, self.y_axis_max])
        plt.grid(which="major", axis='y')
        fig.savefig("fast_data_test.png")
        # plt.show()
        pix_map = QPixmap("fast_data_test.png")
        self.label_12.setPixmap(pix_map)

    def graph_one(self):
        profile_number = "Post"
        voltage_list = []
        time_list = []
        if self.radioButton.isChecked():
            profile_number = "Pre"

        while self.listWidget.currentItem() is None:
            msgbox = QtWidgets.QMessageBox(self)
            msgbox.setText("No item is selected, please select an item.")
            msgbox.exec()
            return

        select_data = self.fast_data_file[self.fast_data_file["Barcode"] == int(self.listWidget.currentItem().text())][profile_number].values[0]
        # if the select data is nan give warning
        if pd.isna(select_data):
            msgbox = QtWidgets.QMessageBox(self)
            msgbox.setText(f"The barcode {self.listWidget.currentItem().text()} you selected has no {profile_number} test data!")
            msgbox.exec()
            return
        select_data_1 = ast.literal_eval(select_data)
        time_check = select_data_1[1][2]

        for item in select_data_1:
            if not voltage_list:
                init_data = self.start_data_file[self.start_data_file["Barcode"] == int(self.listWidget.currentItem().text())][profile_number + '-OCV'].values[0]
                voltage_list.append(init_data)
                time_list.append(-1)
                voltage_list.append(init_data)
                time_list.append(-0.1)
                self.data_set = [(round(float(init_data), 4), -1), (round(float(init_data), 4), -0.1)]
            else:
                voltage_list.append(float(item[0]))
                time_list.append(item[2] - time_check)
                self.data_set.append((round(float(item[0]), 4), round((item[2] - time_check), 4)))
            # time_list.append(item[2] - time_check)
        if self.x_axis_max == 0:
            self.x_axis_max = round(time_list[-1]) + 5

        fig, ax = plt.subplots()
        plt.title("Fast Data plot " + self.test_name[:-4] + " Barcode: " + self.listWidget.currentItem().text(), fontsize=22)
        fig.set_size_inches(23, 18)
        plt.rcParams.update({"font.size": 22})
        plt.tick_params(labelsize=22)
        if self.x_axis_max <= 15:
            plt.xticks(np.arange(-1, self.x_axis_max, 1.0))

        plt.plot(time_list, voltage_list,)
        plt.xlabel('time (S)', fontsize=22)
        plt.ylabel('Voltage (V)', fontsize=22)
        # y-axis lim
        if self.y_axis_min != 0:
            plt.ylim([self.y_axis_min, self.y_axis_max])
        plt.grid(which="major", axis='y')
        fig.savefig("fast_data_test.png")
        # plt.show()
        pix_map = QPixmap("fast_data_test.png")
        self.label_12.setPixmap(pix_map)

    def populate_list_widget(self):
        self.listWidget.clear()
        for item in self.fast_data_file["Barcode"]:
            self.listWidget.addItem(str(item))


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    qt_app = GraphFastDataWithFunction("14891D00.txt")
    qt_app.show()

    def exception_hook(exctype, value, traceback):
        print(exctype, value, traceback)
        sys.excepthook(exctype, value, traceback)
        sys.exit(1)

    sys.excepthook = exception_hook
    app.exec_()