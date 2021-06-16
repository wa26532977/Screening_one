import pandas as pd
import numpy as np
import docx
import os
import sys
from docx.shared import Inches

pd.options.display.max_columns = 999
pd.options.display.max_rows = 999
pd.set_option("display.precision", 6)

def graphPicture_combin(x):
    report_name = x[:-4]
    # import the template
    path_template = os.path.dirname(sys.argv[0]) + r"\\Screening_Template\\" + x
    data_file1 = pd.read_csv(path_template, sep="\t")
    columns_1 = data_file1.loc[0].fillna("")

    doc = docx.Document(os.path.dirname(sys.argv[0]) + r"\\Report_Word\\Doc Template\\Graph_combin.docx")
    # if the cell is tabbed
    if columns_1["Tabbed?"] == "Tabbed":
        doc.add_picture("OCV.png", width=Inches(10.5))
        doc.add_picture("Pre OCV WITHIN SD RANGE.png", width=Inches(10.5))
        doc.add_picture("Pre OCV 95% confidence interval.png", width=Inches(10.5))

        doc.add_picture("OCV 2.png", width=Inches(10.5))
        doc.add_picture("Post OCV WITHIN SD RANGE.png", width=Inches(10.5))
        doc.add_picture("CCV 2.png", width=Inches(10.5))
        doc.add_picture("Post CCV WITHIN SD RANGE.png", width=Inches(10.5))
        doc.add_picture("Post OCV 95% confidence interval.png", width=Inches(10.5))
        doc.add_picture("Post CCV 95% confidence interval.png", width=Inches(10.5))
        doc.add_picture("OCV-CCV 2.png", width=Inches(10.5))
        doc.add_picture("Post OCV-CCV WITHIN SD RANGE.png", width=Inches(10.5))

    elif columns_1["Profile No."] == 2 or columns_1["Section No."] == 2:
        doc.add_picture("OCV.png", width=Inches(10.5))
        doc.add_picture("Pre OCV WITHIN SD RANGE.png", width=Inches(10.5))
        doc.add_picture("CCV.png", width=Inches(10.5))
        doc.add_picture("Pre CCV WITHIN SD RANGE.png", width=Inches(10.5))

        doc.add_picture("Pre OCV 95% confidence interval.png", width=Inches(10.5))
        doc.add_picture("Pre CCV 95% confidence interval.png", width=Inches(10.5))
        doc.add_picture("OCV-CCV.png", width=Inches(10.5))
        doc.add_picture("Pre OCV-CCV WITHIN SD RANGE.png", width=Inches(10.5))

        doc.add_picture("OCV 2.png", width=Inches(10.5))
        doc.add_picture("Post OCV WITHIN SD RANGE.png", width=Inches(10.5))
        doc.add_picture("CCV 2.png", width=Inches(10.5))
        doc.add_picture("Post CCV WITHIN SD RANGE.png", width=Inches(10.5))

        doc.add_picture("Post OCV 95% confidence interval.png", width=Inches(10.5))
        doc.add_picture("Post CCV 95% confidence interval.png", width=Inches(10.5))
        doc.add_picture("OCV-CCV 2.png", width=Inches(10.5))
        doc.add_picture("Post OCV-CCV WITHIN SD RANGE.png", width=Inches(10.5))
    else:
        doc.add_picture("OCV.png", width=Inches(10.5))
        doc.add_picture("Pre OCV WITHIN SD RANGE.png", width=Inches(10.5))
        doc.add_picture("CCV.png", width=Inches(10.5))
        doc.add_picture("Pre CCV WITHIN SD RANGE.png", width=Inches(10.5))
        doc.add_picture("Pre OCV 95% confidence interval.png", width=Inches(10.5))
        doc.add_picture("Pre CCV 95% confidence interval.png", width=Inches(10.5))
        doc.add_picture("OCV-CCV.png", width=Inches(10.5))
        doc.add_picture("Pre OCV-CCV WITHIN SD RANGE.png", width=Inches(10.5))

    doc.save(os.path.dirname(sys.argv[0]) + r"\\Final_Report\\" + report_name + "_Graph.docx")


# for debugging propose
if __name__ == '__main__':
    graphPicture_combin("14462C00.txt")